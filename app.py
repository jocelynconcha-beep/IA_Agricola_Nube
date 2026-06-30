
def mostrar_pdf_bytes(pdf_bytes):
    import base64
    import streamlit.components.v1 as components

    if not pdf_bytes:
        st.warning("No se pudo cargar el PDF.")
        return

    base64_pdf = base64.b64encode(pdf_bytes).decode("utf-8")

    html_pdf = f"""
    <iframe
        src="data:application/pdf;base64,{base64_pdf}"
        width="100%"
        height="850px"
        style="border: 1px solid #d1d5db; border-radius: 10px;"
        type="application/pdf">
    </iframe>
    """

    components.html(html_pdf, height=880, scrolling=True)


import streamlit as st
import os
import zipfile
import sqlite3
import base64
import textwrap
import html
import re
import tempfile
import pandas as pd
from io import BytesIO
from docx import Document
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode

from extractor_pdf import extraer_texto_pdf, extraer_usos_pdf, analizar_texto
from database import (
    crear_tablas,
    guardar_producto,
    actualizar_compatibilidad_producto,
    guardar_usos_producto,
    subir_pdf_storage,
    eliminar_pdf_storage,
    obtener_productos,
    eliminar_producto,
    eliminar_duplicados,
)

crear_tablas()


def aplicar_diseno_responsivo():
    st.markdown("""
    <style>
    /* CONTENEDOR GENERAL */
    .block-container {
        padding-top: 0.35rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }

    .stApp {
        background: #ffffff;
    }

    div[data-testid="stAppViewContainer"],
    section[data-testid="stSidebar"] + div,
    main {
        background: #ffffff !important;
    }

    .app-top-title {
        display: none;
    }

    div[data-testid="stTabs"] [role="tablist"] {
        max-width: 1100px;
        margin: 0 auto 0.5rem auto;
        gap: 0.35rem;
        border-bottom: 0 !important;
    }

    div[data-testid="stTabs"] {
        border-bottom: 0 !important;
    }

    div[data-testid="stTabs"] > div,
    div[data-testid="stTabs"] [role="tabpanel"],
    div[data-testid="stTabs"] [role="tabpanel"] > div {
        border: 0 !important;
        box-shadow: none !important;
        background: #ffffff !important;
    }

    div[data-testid="stTabs"] [data-baseweb="tab-border"] {
        display: none !important;
    }

    div[data-testid="stTabs"] [role="tab"] {
        border-radius: 999px !important;
        padding: 0.5rem 1.1rem !important;
        background: #ffffff !important;
        border: 1px solid #d9e2e8 !important;
        box-shadow: 0 8px 20px rgba(15, 23, 42, 0.05);
        color: #0f172a !important;
        font-weight: 850 !important;
    }

    /* TÍTULOS */
    h1 {
        font-size: 1.35rem !important;
        font-weight: 800 !important;
    }

    h2, h3 {
        font-weight: 750 !important;
    }

    /* BOTONES */
    div.stButton > button,
    div.stDownloadButton > button {
        border-radius: 12px;
        padding: 0.65rem 1rem;
        font-weight: 600;
        border: 1px solid #d0d7de;
    }

    div.stButton > button:hover,
    div.stDownloadButton > button:hover {
        border-color: #58cfa1;
        color: #007a4d;
    }

    /* SELECTBOX Y CAMPOS */
    div[data-baseweb="select"] {
        border-radius: 12px;
    }

    textarea, input {
        border-radius: 12px !important;
    }

    /* ALERTAS */
    div[data-testid="stAlert"] {
        border-radius: 14px;
    }

    /* EXPANDERS */
    details {
        border-radius: 14px !important;
    }

    /* TARJETAS VISUALES */
    .producto-card {
        border: 2px solid #9fe8c8;
        border-radius: 24px;
        padding: 20px;
        margin-bottom: 18px;
        background: #ffffff;
    }

    .producto-card h3 {
        margin-top: 0;
    }

    .modo-terreno-card {
        border: 1px solid #d0d7de;
        border-radius: 18px;
        padding: 18px;
        margin-bottom: 16px;
        background: #ffffff;
        box-shadow: 0 4px 16px rgba(15, 23, 42, 0.06);
    }

    .modo-terreno-card h3 {
        margin: 0 0 6px 0;
        font-size: 1.25rem;
    }

    .modo-terreno-pill {
        display: inline-block;
        padding: 7px 12px;
        margin: 0;
        border-radius: 999px;
        background: #e8f7ef;
        color: #17663d;
        font-size: 0.82rem;
        font-weight: 900;
    }

    .modo-terreno-grid {
        display: grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap: 8px 16px;
        margin-top: 10px;
    }

    .modo-terreno-label {
        color: #57606a;
        font-size: 0.82rem;
        font-weight: 700;
        text-transform: uppercase;
    }

    .modo-terreno-value {
        color: #111827;
        font-size: 0.96rem;
        overflow-wrap: anywhere;
    }

    .modo-terreno-actions div.stButton > button {
        min-height: 78px;
        font-size: 1.08rem;
        border-radius: 20px;
        border: 1px solid #83caa6;
        background: linear-gradient(180deg, #f0fff7 0%, #cfebdc 100%);
        color: #145c38;
        box-shadow: 0 10px 24px rgba(15, 23, 42, 0.12);
    }

    .modo-terreno-action-grid {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 14px;
        margin-top: 16px;
    }

    .modo-terreno-action-card,
.modo-terreno-action-card:link,
.modo-terreno-action-card:visited,
.modo-terreno-action-card:hover,
.modo-terreno-action-card:active {
    text-decoration: none !important;
    color: inherit !important;
}

.modo-terreno-action-card {
        min-height: 86px;
        border-radius: 22px;
        border: 0;
        color: #ffffff;
        box-shadow: 0 10px 24px rgba(15, 23, 42, 0.12);
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 10px;
        text-align: center;
        font-weight: 850;
        font-size: 1.05rem;
    }

    .modo-terreno-action-card.pdf {
        background: linear-gradient(180deg, #35b86f 0%, #168a4f 100%);
    }

    .modo-terreno-action-card.campo {
        background: linear-gradient(180deg, #28b8a8 0%, #087f79 100%);
    }

    .modo-terreno-action-card.db {
        background: linear-gradient(180deg, #3b82f6 0%, #1d4ed8 100%);
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-app-shell-marker) {
        max-width: 1100px;
        margin: 0 auto;
        padding: 18px 22px;
        border: 0 !important;
        border-radius: 0;
        background: #ffffff;
        box-shadow: none !important;
    }

    div[data-testid="stVerticalBlockBorderWrapper"]:has(.modo-terreno-app-shell-marker) {
        border: 0 !important;
        box-shadow: none !important;
        background: transparent !important;
        padding: 0 !important;
        overflow: hidden;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-app-shell-marker) h3 {
        margin: 0.45rem 0 0.25rem 0 !important;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-app-shell-marker) div[data-testid="stButton"] > button {
        min-height: 42px;
        padding: 0.45rem 0.8rem;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-app-shell-marker) div[data-testid="stTextInput"] input {
        min-height: 42px;
        padding-top: 0.45rem;
        padding-bottom: 0.45rem;
    }

    .modo-terreno-app-shell-marker {
        display: none;
    }

    .modo-terreno-action-icon {
        width: 28px;
        height: 28px;
        border-radius: 10px;
        background: rgba(255, 255, 255, 0.18);
        display: flex;
        align-items: center;
        justify-content: center;
        position: relative;
        flex: 0 0 auto;
    }

    .modo-terreno-action-icon::before,
    .modo-terreno-action-icon::after {
        content: "";
        position: absolute;
        display: block;
    }

    .modo-terreno-action-icon.pdf::before {
        width: 13px;
        height: 16px;
        border: 2px solid #ffffff;
        border-radius: 2px;
    }

    .modo-terreno-action-icon.pdf::after {
        width: 6px;
        height: 6px;
        border-top: 2px solid #ffffff;
        border-right: 2px solid #ffffff;
        right: 7px;
        top: 6px;
    }

    .modo-terreno-action-icon.campo::before {
        width: 15px;
        height: 17px;
        border: 2px solid #ffffff;
        border-radius: 3px;
    }

    .modo-terreno-action-icon.campo::after {
        width: 10px;
        height: 2px;
        background: #ffffff;
        box-shadow: 0 5px 0 #ffffff, 0 10px 0 #ffffff;
    }

    .modo-terreno-action-icon.db::before {
        width: 17px;
        height: 9px;
        border: 2px solid #ffffff;
        border-radius: 50%;
        top: 6px;
    }

    .modo-terreno-action-icon.db::after {
        width: 17px;
        height: 12px;
        border-left: 2px solid #ffffff;
        border-right: 2px solid #ffffff;
        border-bottom: 2px solid #ffffff;
        border-radius: 0 0 9px 9px;
        bottom: 5px;
    }

    .modo-terreno-selector-card {
        min-height: 82px;
        border-top: 6px solid #58cfa1;
        border-radius: 18px;
        padding: 14px 14px 10px 14px;
        margin: -4px -4px 6px -4px;
        background: linear-gradient(180deg, #ffffff 0%, #f7fbff 100%);
        display: grid;
        grid-template-columns: 50px 1fr;
        align-items: center;
        gap: 12px;
        text-align: left;
    }

    .modo-terreno-selector-card.selector-cultivo {
        border-top-color: #12965b;
    }

    .modo-terreno-selector-card.selector-plaga {
        border-top-color: #f97316;
    }

    .modo-terreno-selector-card.selector-enfermedad {
        border-top-color: #0b67c2;
    }

    .modo-terreno-selector-card.selector-maleza {
        border-top-color: #7c5ce6;
    }

    .modo-terreno-selector-icon {
        width: 44px;
        height: 44px;
        border-radius: 0;
        display: flex;
        align-items: center;
        justify-content: center;
        background: transparent;
        color: #17663d;
        position: relative;
        margin: 0;
        font-size: 2rem;
        line-height: 1;
    }

    .modo-terreno-selector-emoji {
        display: block;
        line-height: 1;
    }

    .modo-terreno-selector-icon svg {
        width: 38px;
        height: 38px;
        display: block;
        stroke: currentColor;
        fill: none;
        stroke-width: 2.4;
        stroke-linecap: round;
        stroke-linejoin: round;
    }

    .modo-terreno-selector-icon::before {
        content: "";
        width: 38px;
        height: 38px;
        display: none;
        background: currentColor;
        -webkit-mask-position: center;
        -webkit-mask-repeat: no-repeat;
        -webkit-mask-size: contain;
        mask-position: center;
        mask-repeat: no-repeat;
        mask-size: contain;
    }

    .selector-cultivo .modo-terreno-selector-icon::before {
        -webkit-mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M11 20A7 7 0 0 1 4 13c0-5 5-9 15-9 0 10-4 15-9 15Z' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3Cpath d='M4 20c4-6 8-9 15-12' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3C/svg%3E");
        mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M11 20A7 7 0 0 1 4 13c0-5 5-9 15-9 0 10-4 15-9 15Z' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3Cpath d='M4 20c4-6 8-9 15-12' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3C/svg%3E");
    }

    .selector-plaga .modo-terreno-selector-icon::before {
        -webkit-mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M8 2l1.88 1.88' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3Cpath d='M14.12 3.88 16 2' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3Cpath d='M9 7.13v-1a3 3 0 0 1 6 0v1' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3Cpath d='M12 20c-3.3 0-6-2.7-6-6v-3a6 6 0 0 1 12 0v3c0 3.3-2.7 6-6 6Z' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3Cpath d='M12 20v-9M6 13H2M22 13h-4M6.7 17 3 19M17.3 17 21 19M6.7 9 3 7M17.3 9 21 7' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3C/svg%3E");
        mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M8 2l1.88 1.88' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3Cpath d='M14.12 3.88 16 2' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3Cpath d='M9 7.13v-1a3 3 0 0 1 6 0v1' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3Cpath d='M12 20c-3.3 0-6-2.7-6-6v-3a6 6 0 0 1 12 0v3c0 3.3-2.7 6-6 6Z' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3Cpath d='M12 20v-9M6 13H2M22 13h-4M6.7 17 3 19M17.3 17 21 19M6.7 9 3 7M17.3 9 21 7' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round'/%3E%3C/svg%3E");
    }

    .selector-enfermedad .modo-terreno-selector-icon::before {
        -webkit-mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M20 13c0 5-3.5 7.5-8 9-4.5-1.5-8-4-8-9V5l8-3 8 3v8Z' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3Cpath d='m9 12 2 2 4-5' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3C/svg%3E");
        mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M20 13c0 5-3.5 7.5-8 9-4.5-1.5-8-4-8-9V5l8-3 8 3v8Z' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3Cpath d='m9 12 2 2 4-5' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3C/svg%3E");
    }

    .selector-maleza .modo-terreno-selector-icon::before {
        -webkit-mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M12 22V9M12 13c-3.5 0-6-2.5-6-6 3.5 0 6 2.5 6 6ZM12 15c3.5 0 6-2.5 6-6-3.5 0-6 2.5-6 6ZM5 22h14' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3C/svg%3E");
        mask-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 24 24' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M12 22V9M12 13c-3.5 0-6-2.5-6-6 3.5 0 6 2.5 6 6ZM12 15c3.5 0 6-2.5 6-6-3.5 0-6 2.5-6 6ZM5 22h14' fill='none' stroke='black' stroke-width='2.4' stroke-linecap='round' stroke-linejoin='round'/%3E%3C/svg%3E");
    }

    .modo-terreno-selector-card.selector-cultivo .modo-terreno-selector-icon {
        color: #149358;
        background: transparent;
    }

    .modo-terreno-selector-card.selector-plaga .modo-terreno-selector-icon {
        color: #f97316;
        background: transparent;
    }

    .modo-terreno-selector-card.selector-enfermedad .modo-terreno-selector-icon {
        color: #0b67c2;
        background: transparent;
    }

    .modo-terreno-selector-card.selector-maleza .modo-terreno-selector-icon {
        color: #7c5ce6;
        background: transparent;
    }

    .modo-terreno-filter-tools {
        display: flex;
        justify-content: flex-end;
        align-items: center;
        gap: 10px;
        margin: 6px 0 8px 0;
    }

    .modo-terreno-selector-title {
        font-weight: 800;
        font-size: 1.12rem;
        color: #111827;
    }

    .modo-terreno-selector-value {
        color: #57606a;
        font-size: 0.82rem;
        margin-top: 2px;
        overflow-wrap: anywhere;
    }

    .modo-terreno-products-head {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 12px;
        padding: 0 4px 10px 4px;
        border-bottom: 1px solid #edf1f5;
        margin-bottom: 8px;
        color: #0f172a;
    }

    .modo-terreno-products-title {
        display: flex;
        align-items: center;
        gap: 8px;
        font-weight: 850;
        font-size: 0.98rem;
    }

    .modo-terreno-products-title::before {
        content: "";
        width: 18px;
        height: 22px;
        border: 2px solid #0b67c2;
        border-radius: 4px;
        display: inline-block;
    }

    .modo-terreno-products-hint {
        color: #57606a;
        font-size: 0.82rem;
    }

    .modo-terreno-lista {
        border: 1px solid #d0d7de;
        border-radius: 22px;
        background: #ffffff;
        padding: 10px;
        box-shadow: 0 10px 28px rgba(15, 23, 42, 0.07);
        margin: 10px 0 20px 0;
    }

    div[data-testid="stVerticalBlockBorderWrapper"]:has(.modo-terreno-lista-botones-marker) {
        border-radius: 20px;
        border-color: #d9e2e8;
        box-shadow: 0 10px 26px rgba(15, 23, 42, 0.06);
        background: #ffffff;
    }

    div[data-testid="stVerticalBlockBorderWrapper"]:has(.modo-terreno-lista-botones-marker) > div {
        padding: 12px 16px !important;
    }

    .modo-terreno-product-row-html {
        display: grid;
        grid-template-columns: 46px minmax(190px, 1.35fr) minmax(118px, 0.55fr) minmax(120px, 0.55fr) minmax(106px, 0.45fr);
        align-items: center;
        gap: 14px;
        min-height: 58px;
        border: 0;
        border-bottom: 1px solid #edf1f5;
        border-radius: 0;
        padding: 8px 4px;
        background: #ffffff;
        margin: 0;
    }

    .modo-terreno-product-row-html:hover {
        background: #f8fcfa;
        box-shadow: none;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-row-html) {
        gap: 0.45rem;
        align-items: center;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-name-button-marker) {
        gap: 0.45rem;
        align-items: center;
        border-bottom: 1px solid #edf1f5;
        padding: 8px 4px;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-name-button-marker) div[data-testid="stVerticalBlockBorderWrapper"] {
        border: 0;
        box-shadow: none;
        background: transparent;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-name-button-marker) div.stButton > button {
        border: 0;
        background: transparent;
        box-shadow: none;
        padding: 0;
        min-height: auto;
        color: #0f172a;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-name-button-marker) div.stButton > button:hover {
        border: 0;
        background: transparent;
        color: #17663d;
        box-shadow: none;
        text-decoration: underline;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-full-button-marker) div.stButton > button {
        width: 100%;
        min-height: 74px;
        justify-content: flex-start;
        text-align: left;
        border: 1px solid #edf1f5;
        border-radius: 16px;
        padding: 10px 16px;
        margin: 0 0 8px 0;
        background: #ffffff;
        color: #0f172a;
        box-shadow: none;
        font-weight: 800;
        line-height: 1.35;
        white-space: pre-wrap;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-full-button-marker) div.stButton > button:hover {
        border-color: #58cfa1;
        background: #f8fcfa;
        color: #0f172a;
        box-shadow: 0 8px 18px rgba(15, 23, 42, 0.08);
        transform: translateY(-1px);
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-full-button-marker) div.stButton > button:focus {
        border-color: #12965b;
        box-shadow: 0 0 0 3px rgba(18, 150, 91, 0.12);
    }

    .modo-terreno-product-full-button-marker {
        display: none;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-row-click-marker) {
        position: relative;
        margin-bottom: 8px;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-row-click-marker) div.stButton {
        position: absolute;
        inset: 0;
        z-index: 5;
        margin: 0;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-row-click-marker) div.stButton > button {
        width: 100%;
        height: 100%;
        min-height: 74px;
        opacity: 0;
        cursor: pointer;
        border: 0;
        padding: 0;
    }

    .modo-terreno-product-row-click-marker {
        display: none;
    }

    .modo-terreno-product-row-pretty {
        display: grid;
        grid-template-columns: 58px minmax(230px, 2fr) minmax(130px, 0.8fr) minmax(170px, 1fr) minmax(130px, 0.8fr) 42px;
        align-items: center;
        gap: 16px;
        min-height: 74px;
        border-bottom: 1px solid #edf1f5;
        background: #ffffff;
        padding: 10px 6px;
        transition: background 0.15s ease, box-shadow 0.15s ease;
    }

    .modo-terreno-product-row-pretty:hover {
        background: #f8fcfa;
        box-shadow: inset 0 0 0 1px #d9f1e2;
    }

    .modo-terreno-product-name-box {
        display: block;
        border: 0;
        border-radius: 0;
        padding: 0;
        color: #0f172a;
        font-size: 1rem;
        font-weight: 900;
        line-height: 1.15;
        background: transparent;
        margin-bottom: 6px;
        letter-spacing: 0;
    }

    .modo-terreno-product-row-pretty-arrow {
        width: 40px;
        height: 40px;
        border: 0;
        border-radius: 999px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: #17663d;
        font-size: 1.45rem;
        font-weight: 950;
        background: #e9f8f0;
    }

    .modo-terreno-product-row-link {
        display: grid;
        grid-template-columns: 48px minmax(245px, 2.2fr) minmax(120px, 0.72fr) minmax(190px, 1.1fr) minmax(118px, 0.72fr);
        align-items: center;
        gap: 14px;
        min-height: 78px;
        border: 1px solid #dde7ed;
        border-radius: 17px;
        background: #fbfefd;
        padding: 11px 14px;
        margin: 6px 0;
        color: inherit;
        text-decoration: none;
        cursor: pointer;
        transition: background 0.15s ease, box-shadow 0.15s ease;
    }

    .modo-terreno-product-row-link:hover {
        background: #f8fcfa;
        box-shadow: inset 0 0 0 1px #cfeedd, 0 8px 18px rgba(15, 23, 42, 0.06);
        color: inherit;
        text-decoration: none;
    }

    .modo-terreno-product-row-link:hover .modo-terreno-product-name-box {
        color: #17663d;
        box-shadow: none;
    }

    div[data-testid="stVerticalBlock"]:has(> div[data-testid="stElementContainer"] .modo-terreno-product-click-row-marker) {
        position: relative;
        min-height: 78px;
    }

    div[data-testid="stVerticalBlock"]:has(> div[data-testid="stElementContainer"] .modo-terreno-product-click-row-marker) > div[data-testid="stElementContainer"]:has(div[data-testid="stButton"]) {
        position: absolute;
        inset: 0;
        z-index: 8;
        margin: 0;
    }

    div[data-testid="stVerticalBlock"]:has(> div[data-testid="stElementContainer"] .modo-terreno-product-click-row-marker) > div[data-testid="stElementContainer"] div[data-testid="stButton"] > button {
        width: 100%;
        height: 100%;
        min-height: 78px;
        opacity: 0;
        cursor: pointer;
        border: 0;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-row-line-marker) {
        align-items: center;
        border-bottom: 1px solid #edf1f5;
        padding: 10px 6px;
        gap: 0.55rem;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-row-line-marker):hover {
        background: #f8fcfa;
        box-shadow: inset 0 0 0 1px #d9f1e2;
    }

    .modo-terreno-product-row-line-marker,
    .modo-terreno-product-name-real-button-marker {
        display: none;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-name-real-button-marker) div.stButton > button {
        display: inline-flex;
        width: auto;
        min-height: auto;
        border: 1px solid #d6dde5;
        border-radius: 12px;
        padding: 9px 14px;
        background: #ffffff;
        color: #0f172a;
        box-shadow: none;
        font-size: 0.98rem;
        font-weight: 850;
        line-height: 1.2;
        text-align: left;
        margin-bottom: 0;
    }

    div[data-testid="stVerticalBlock"]:has(.modo-terreno-product-name-real-button-marker) div.stButton > button:hover {
        border-color: #58cfa1;
        background: #ffffff;
        color: #17663d;
        box-shadow: 0 6px 14px rgba(15, 23, 42, 0.08);
        text-decoration: none;
    }

    .modo-terreno-product-name-button-marker {
        display: none;
    }

    .modo-terreno-product-icon-cell {
        display: flex;
        align-items: center;
        justify-content: center;
        min-height: 48px;
    }

    .modo-terreno-product-name-meta {
        color: #57606a;
        font-size: 0.93rem;
        font-weight: 650;
        margin-top: 0;
        line-height: 1.25;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-row-html) div.stButton > button {
        width: 42px;
        height: 42px;
        min-height: 42px;
        border-radius: 999px;
        padding: 0;
        border: 0;
        background: #e8f7ef;
        color: #17663d;
        font-size: 1.35rem;
        font-weight: 950;
        box-shadow: none;
        margin-top: 11px;
    }

    div[data-testid="stHorizontalBlock"]:has(.modo-terreno-product-row-html) div.stButton > button:hover {
        background: #c9efd9;
        color: #0f7a47;
        border: 0;
        box-shadow: 0 8px 18px rgba(22, 122, 71, 0.18);
    }

    .modo-terreno-lista-item {
        display: grid;
        grid-template-columns: 38px minmax(190px, 1.2fr) minmax(140px, 0.9fr) minmax(130px, 0.8fr) minmax(98px, 0.6fr) 30px;
        gap: 10px;
        align-items: center;
        border: 1px solid #e6ebef;
        border-radius: 15px;
        padding: 9px 12px;
        margin-bottom: 8px;
        background: #fbfdfc;
        min-height: 70px;
        cursor: pointer;
        text-decoration: none;
        color: inherit;
        transition: transform 0.15s ease, border-color 0.15s ease, box-shadow 0.15s ease;
    }

    .modo-terreno-lista-item:last-child {
        margin-bottom: 0;
    }

    .modo-terreno-lista-item:hover {
        transform: translateY(-2px);
        border-color: #58cfa1;
        box-shadow: 0 10px 24px rgba(15, 23, 42, 0.1);
    }

    .modo-terreno-lista-nombre {
        font-weight: 800;
        color: #0f172a;
        font-size: 0.98rem;
        margin-bottom: 2px;
    }

    .modo-terreno-lista-meta {
        color: #57606a;
        font-size: 0.93rem;
        line-height: 1.35;
        overflow-wrap: anywhere;
    }

    .modo-terreno-lista-meta strong {
        color: #515d6b;
        font-size: 0.92rem;
        font-weight: 900;
    }

    .modo-terreno-lista-row {
        display: contents;
    }

    .modo-terreno-product-icon {
        width: 44px;
        height: 44px;
        border-radius: 15px;
        background: #e8f7ef;
        color: #17663d;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 900;
    }

    .modo-terreno-product-icon::before {
        content: "";
        width: 18px;
        height: 22px;
        border: 2px solid currentColor;
        border-radius: 4px 4px 7px 7px;
        display: block;
    }

    .modo-terreno-product-icon.fungicida {
        background: #e9f3ff;
        color: #0b67c2;
    }

    .modo-terreno-product-icon.herbicida {
        background: #f1edff;
        color: #7c5ce6;
    }

    .modo-terreno-product-icon.herbicida::before {
        border: 0;
        width: 20px;
        height: 20px;
        background: currentColor;
        border-radius: 70% 10% 70% 10%;
        transform: rotate(-35deg);
    }

    .modo-terreno-arrow {
        width: 28px;
        height: 28px;
        border-radius: 999px;
        background: #e8f7ef;
        color: #17663d;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 900;
        font-size: 1.15rem;
    }

    .modo-terreno-bee {
        display: inline-block;
        border-radius: 999px;
        padding: 7px 12px;
        font-size: 0.82rem;
        font-weight: 900;
        vertical-align: middle;
    }

    .modo-terreno-bee.alta {
        background: #ffe8e8;
        color: #a61b1b;
    }

    .modo-terreno-bee.media {
        background: #fff3cd;
        color: #8a5a00;
    }

    .modo-terreno-bee.baja {
        background: #e8f7ef;
        color: #17663d;
    }

    .modo-terreno-ficha {
        border: 1px solid #c9d8d0;
        border-radius: 24px;
        padding: 24px;
        background: #ffffff;
        box-shadow: 0 14px 34px rgba(15, 23, 42, 0.08);
        margin-top: 12px;
    }

    .modo-terreno-ficha-header {
        display: flex;
        justify-content: space-between;
        gap: 16px;
        align-items: flex-start;
        border-bottom: 1px solid #eaeef2;
        padding-bottom: 14px;
        margin-bottom: 16px;
    }

    .modo-terreno-ficha-title {
        font-size: 2rem;
        font-weight: 850;
        color: #0f172a;
        margin: 0;
        overflow-wrap: anywhere;
    }

    .modo-terreno-ficha-grid {
        display: grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap: 12px;
    }

    .modo-terreno-ficha-section {
        border: 1px solid #e5e7eb;
        border-radius: 16px;
        padding: 14px;
        background: #fbfcfd;
        min-height: 104px;
    }

    .modo-terreno-ficha-section.full {
        grid-column: 1 / -1;
        background: #f3fbf6;
        border-color: #bfe7cf;
    }

    .modo-terreno-ficha-section-title {
        display: flex;
        align-items: center;
        gap: 8px;
        color: #17663d;
        font-weight: 850;
        margin-bottom: 8px;
    }

    .modo-terreno-ficha-section-value {
        color: #111827;
        overflow-wrap: anywhere;
        line-height: 1.42;
    }

    .modo-terreno-info-box {
        width: min(920px, 100%);
        margin: 2px auto 8px auto;
        border: 2px dashed #a9d8ff;
        border-radius: 16px;
        background: #f8fcff;
        padding: 8px 14px;
        text-align: center;
        box-shadow: 0 8px 22px rgba(15, 23, 42, 0.04);
    }

    .modo-terreno-info-title {
        font-size: 0.98rem;
        font-weight: 850;
        color: #0f172a;
        margin-bottom: 2px;
    }

    .modo-terreno-info-subtitle {
        color: #57606a;
        font-size: 0.84rem;
    }

    .modo-terreno-compat-wrapper {
        border: 1px solid #d0d7de;
        border-radius: 24px;
        background: #ffffff;
        padding: 12px;
        box-shadow: 0 10px 28px rgba(15, 23, 42, 0.07);
        margin-top: 6px;
    }

    .modo-terreno-section-title {
        font-weight: 850;
        color: #0f172a;
        font-size: 1.18rem;
        margin: 18px 0 8px 0;
    }

    .modo-terreno-count {
        color: #57606a;
        font-size: 0.9rem;
        margin: 4px 0 10px 0;
    }

    div[data-testid="stVerticalBlockBorderWrapper"] {
        border-radius: 24px;
        border-color: #d9e2e8;
        box-shadow: 0 10px 26px rgba(15, 23, 42, 0.06);
        background: #ffffff;
    }

    div[data-testid="stVerticalBlock"] > div:has(.modo-terreno-selector-card) {
        gap: 0.25rem;
    }

    h3 {
        margin-top: 0.65rem !important;
        margin-bottom: 0.35rem !important;
    }

    .modo-terreno-compat-title {
        font-size: 1.12rem;
        font-weight: 850;
        color: #0f172a;
        margin-bottom: 8px;
    }

    .modo-terreno-compat-slots {
        display: grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap: 8px;
        margin-bottom: 8px;
    }

    .modo-terreno-compat-slot {
        border: 1px dashed #9fd8bd;
        border-radius: 16px;
        background: #f6fffa;
        padding: 10px;
        color: #17663d;
        font-weight: 800;
        min-height: 46px;
        display: flex;
        align-items: center;
        justify-content: center;
        text-align: center;
        overflow-wrap: anywhere;
    }

    .modo-terreno-compat-result {
        border: 1px solid #bfe7cf;
        border-radius: 18px;
        background: #edf9f2;
        padding: 10px 12px;
        color: #145c38;
        font-weight: 700;
    }

    .base-app-shell {
        max-width: 1100px;
        margin: 0 auto 1.1rem auto;
        padding: 0;
    }

    .base-hero {
        min-height: 138px;
        border-radius: 18px;
        background: radial-gradient(circle at 92% 10%, rgba(255,255,255,0.16) 0 10%, transparent 11%),
                    linear-gradient(135deg, #0b4baa 0%, #0754b8 54%, #0a65c8 100%);
        color: #ffffff;
        padding: 28px 40px;
        display: flex;
        align-items: center;
        gap: 26px;
        box-shadow: 0 12px 26px rgba(11, 75, 170, 0.25);
        margin-bottom: 30px;
        overflow: hidden;
        position: relative;
    }

    .base-hero-icon {
        width: 76px;
        height: 76px;
        border-radius: 18px;
        background: rgba(255, 255, 255, 0.14);
        position: relative;
        flex: 0 0 auto;
    }

    .base-hero-icon::before,
    .base-hero-icon::after {
        content: "";
        position: absolute;
        left: 22px;
        width: 32px;
        border: 4px solid #ffffff;
    }

    .base-hero-icon::before {
        top: 20px;
        height: 18px;
        border-radius: 50%;
    }

    .base-hero-icon::after {
        top: 31px;
        height: 25px;
        border-top: 0;
        border-radius: 0 0 18px 18px;
    }

    .base-hero h2 {
        margin: 0 0 8px 0;
        font-size: 2.55rem;
        line-height: 1;
        color: #ffffff;
    }

    .base-hero p {
        margin: 0;
        font-size: 1.05rem;
        color: rgba(255,255,255,0.92);
    }

    .base-table-card {
        border: 1px solid #d7dde5;
        border-radius: 18px;
        background: #ffffff;
        box-shadow: 0 12px 30px rgba(15, 23, 42, 0.08);
        overflow: hidden;
        margin-bottom: 24px;
    }

    .base-table-head,
    .base-table-row {
        display: grid;
        grid-template-columns: minmax(240px, 1.25fr) minmax(180px, 1fr) minmax(120px, 0.72fr) minmax(180px, 1fr);
        align-items: center;
    }

    .base-table-head {
        min-height: 70px;
        color: #0f172a;
        font-weight: 850;
        font-size: 0.98rem;
        border-bottom: 1px solid #dfe5eb;
    }

    .base-table-cell {
        padding: 0 26px;
        min-width: 0;
    }

    .base-table-row {
        min-height: 70px;
        border-bottom: 1px solid #edf1f5;
        color: #334155;
    }

    .base-table-row:last-child {
        border-bottom: 0;
    }

    .base-product-cell {
        display: flex;
        align-items: center;
        gap: 14px;
        font-weight: 850;
        color: #0f172a;
    }

    .base-product-icon {
        width: 42px;
        height: 42px;
        border-radius: 10px;
        border: 1px solid #bae6c8;
        background: #ecfdf3;
        position: relative;
        flex: 0 0 auto;
    }

    .base-product-icon::before {
        content: "";
        position: absolute;
        inset: 9px 13px 8px 13px;
        border: 3px solid #16994f;
        border-radius: 4px 4px 7px 7px;
    }

    .base-product-icon.fungicida {
        border-color: #ffd68a;
        background: #fff7e5;
    }

    .base-product-icon.fungicida::before {
        border-color: #f59e0b;
    }

    .base-product-icon.herbicida {
        border-color: #c4b5fd;
        background: #f3efff;
    }

    .base-product-icon.herbicida::before {
        border-color: #7c3aed;
        border-radius: 70% 10% 70% 10%;
        transform: rotate(-35deg);
        border-width: 0;
        background: #7c3aed;
    }

    .base-actions-grid {
        display: grid;
        grid-template-columns: minmax(0, 1.25fr) minmax(280px, 0.85fr);
        gap: 28px;
        margin-bottom: 22px;
    }

    .base-select-card,
    .base-pdf-stack-card {
        border: 1px solid #d7dde5;
        border-radius: 18px;
        background: #ffffff;
        box-shadow: 0 12px 30px rgba(15, 23, 42, 0.08);
        padding: 26px;
    }

    .base-search-box {
        margin-top: 22px;
        border: 1px solid #f7c65d;
        border-radius: 14px;
        background: #fff9e8;
        min-height: 78px;
        display: flex;
        align-items: center;
        gap: 16px;
        padding: 0 20px;
        color: #85611a;
        font-weight: 700;
    }

    .base-pdf-button {
        min-height: 118px;
        border-radius: 18px;
        background: linear-gradient(135deg, #0b4baa 0%, #065fbd 100%);
        color: #ffffff;
        display: flex;
        align-items: center;
        gap: 22px;
        padding: 0 28px;
        margin-bottom: 24px;
        font-size: 1.25rem;
        font-weight: 850;
        box-shadow: 0 12px 24px rgba(11, 75, 170, 0.2);
    }

    .base-pdf-button:last-child {
        margin-bottom: 0;
    }

    .base-pdf-icon {
        width: 52px;
        height: 62px;
        border: 3px solid #ffffff;
        border-radius: 9px;
        position: relative;
        flex: 0 0 auto;
    }

    .base-pdf-icon::after {
        content: "PDF";
        position: absolute;
        left: 7px;
        bottom: 10px;
        font-size: 0.94rem;
        font-weight: 900;
    }

    /* TABLET */
    @media screen and (max-width: 900px) {
        .block-container {
            padding-left: 1rem;
            padding-right: 1rem;
            max-width: 100%;
        }

        h1 {
            font-size: 2.3rem !important;
        }

        h2 {
            font-size: 1.8rem !important;
        }

        h3 {
            font-size: 1.35rem !important;
        }

        div.stButton > button,
        div.stDownloadButton > button {
            width: 100%;
            margin-bottom: 0.4rem;
        }
    }

    /* TELÉFONO */
    @media screen and (max-width: 600px) {
        .block-container {
            padding-left: 0.7rem;
            padding-right: 0.7rem;
            padding-top: 1rem;
        }

        h1 {
            font-size: 2rem !important;
        }

        h2 {
            font-size: 1.55rem !important;
        }

        h3 {
            font-size: 1.2rem !important;
        }

        p, li, label, div {
            font-size: 0.96rem;
        }

        div.stButton > button,
        div.stDownloadButton > button {
            width: 100%;
            min-height: 44px;
        }

        iframe {
            max-width: 100% !important;
        }

        [data-testid="stHorizontalBlock"] {
            gap: 0.4rem;
        }

        .modo-terreno-grid {
            grid-template-columns: 1fr;
        }

        .modo-terreno-lista-item,
        .modo-terreno-ficha-grid {
            grid-template-columns: 1fr;
        }

        .modo-terreno-lista,
        .modo-terreno-lista-row {
            grid-template-columns: 1fr;
        }

        .modo-terreno-compat-slots {
            grid-template-columns: 1fr;
        }

        .modo-terreno-lista-item {
            grid-template-columns: 36px 1fr 30px;
        }

        .modo-terreno-lista-row {
            display: grid;
            grid-template-columns: 1fr;
            grid-column: 2 / -1;
        }

        .modo-terreno-action-grid {
            grid-template-columns: 1fr;
        }

        .base-hero {
            padding: 22px;
            gap: 16px;
        }

        .base-hero h2 {
            font-size: 1.9rem;
        }

        .base-table-head {
            display: none;
        }

        .base-table-row {
            grid-template-columns: 1fr;
            gap: 5px;
            padding: 14px 0;
        }

        .base-table-cell {
            padding: 0 18px;
        }

        .base-actions-grid {
            grid-template-columns: 1fr;
        }

        div[data-testid="stVerticalBlockBorderWrapper"]:has(.modo-terreno-app-shell-marker) {
            border: 0 !important;
            box-shadow: none !important;
            background: transparent !important;
            padding: 0 !important;
        }

        div[data-testid="stVerticalBlock"]:has(.modo-terreno-app-shell-marker) {
            padding: 14px 12px;
            border-radius: 0;
            border: 0 !important;
            box-shadow: none !important;
        }

        .modo-terreno-ficha-header {
            display: block;
        }

        .modo-terreno-ficha-title {
            font-size: 1.35rem;
        }
    }
    </style>
    """, unsafe_allow_html=True)





# Ajustes finales para teléfonos
st.markdown(
    """
    <style>
    @media screen and (max-width: 600px) {

        html,
        body,
        [data-testid="stAppViewContainer"],
        [data-testid="stAppViewContainer"] > .main {
            max-width: 100% !important;
            overflow-x: hidden !important;
        }

        .block-container {
            max-width: 100% !important;
            padding: 0.75rem 0.65rem 6rem 0.65rem !important;
        }

        /* Ocultar elementos superiores que quitan espacio */
        [data-testid="stToolbar"],
        [data-testid="stDecoration"],
        .stDeployButton {
            display: none !important;
        }

        header[data-testid="stHeader"] {
            height: 0 !important;
            min-height: 0 !important;
        }

        /* Tarjetas de Cultivo, Plaga, Enfermedad y Maleza */
        .modo-terreno-selector-card {
            min-height: auto !important;
            padding: 12px 12px 10px 12px !important;
            margin: 0 0 10px 0 !important;
            border-radius: 17px !important;
        }

        .modo-terreno-selector-icon {
            width: 52px !important;
            height: 52px !important;
            min-width: 52px !important;
        }

        .modo-terreno-selector-emoji {
            font-size: 1.75rem !important;
        }

        .modo-terreno-selector-title {
            font-size: 1.05rem !important;
            line-height: 1.15 !important;
        }

        .modo-terreno-selector-value {
            font-size: 0.82rem !important;
            line-height: 1.3 !important;
        }

        /* Selectores */
        div[data-baseweb="select"] > div {
            min-height: 48px !important;
            border-radius: 12px !important;
        }

        div[data-baseweb="select"] span {
            font-size: 0.95rem !important;
        }

        /* Encabezado de productos */
        .modo-terreno-products-head {
            display: grid !important;
            grid-template-columns: 1fr !important;
            gap: 5px !important;
            padding: 4px 2px 10px 2px !important;
        }

        .modo-terreno-products-title {
            font-size: 1rem !important;
            line-height: 1.35 !important;
        }

        .modo-terreno-products-hint {
            font-size: 0.8rem !important;
            line-height: 1.35 !important;
        }

        /* Contenedor de lista */
        .modo-terreno-lista,
        div[data-testid="stVerticalBlockBorderWrapper"]:has(
            .modo-terreno-lista-botones-marker
        ) {
            width: 100% !important;
            max-width: 100% !important;
            min-width: 0 !important;
            margin-left: 0 !important;
            margin-right: 0 !important;
            box-sizing: border-box !important;
            border-radius: 16px !important;
        }

        div[data-testid="stVerticalBlockBorderWrapper"]:has(
            .modo-terreno-lista-botones-marker
        ) > div {
            padding: 10px !important;
        }

        /* Filas de productos */
        .modo-terreno-product-row-link,
        .modo-terreno-product-row-pretty,
        .modo-terreno-product-row-html {
            width: 100% !important;
            max-width: 100% !important;
            min-width: 0 !important;
            box-sizing: border-box !important;
            overflow: hidden !important;
        }

        .modo-terreno-product-row-link {
            grid-template-columns: 48px minmax(0, 1fr) !important;
            gap: 10px !important;
            min-height: 72px !important;
            padding: 10px !important;
            border-radius: 15px !important;
        }

        .modo-terreno-product-row-pretty {
            grid-template-columns: 48px minmax(0, 1fr) !important;
            gap: 10px !important;
            min-height: 70px !important;
            padding: 10px 6px !important;
        }

        .modo-terreno-product-row-html {
            grid-template-columns: 42px minmax(0, 1fr) !important;
            gap: 9px !important;
        }

        .modo-terreno-product-name-box {
            min-width: 0 !important;
            max-width: 100% !important;
            white-space: normal !important;
            overflow-wrap: anywhere !important;
            word-break: break-word !important;
            font-size: 0.95rem !important;
            line-height: 1.2 !important;
            margin-bottom: 4px !important;
        }

        div.stButton > button {
            white-space: normal !important;
            overflow-wrap: anywhere !important;
            word-break: break-word !important;
        }

        /* Compatibilidad */
        .modo-terreno-compat-wrapper {
            width: 100% !important;
            max-width: 100% !important;
            box-sizing: border-box !important;
            padding: 14px 12px !important;
            border-radius: 18px !important;
        }

        .modo-terreno-compat-title {
            font-size: 1.25rem !important;
            margin-bottom: 10px !important;
        }

        .modo-terreno-compat-slots {
            grid-template-columns: 1fr !important;
            gap: 8px !important;
        }

        .modo-terreno-compat-slot {
            min-width: 0 !important;
            min-height: 58px !important;
            padding: 10px !important;
            font-size: 0.95rem !important;
            line-height: 1.25 !important;
            white-space: normal !important;
            overflow-wrap: anywhere !important;
            word-break: break-word !important;
        }

        .modo-terreno-compat-result {
            width: 100% !important;
            max-width: 100% !important;
            box-sizing: border-box !important;
            padding: 14px !important;
            font-size: 0.92rem !important;
            line-height: 1.45 !important;
            overflow-wrap: anywhere !important;
        }

        /* Botones inferiores */
        .modo-terreno-action-grid {
            grid-template-columns: 1fr !important;
            gap: 10px !important;
            margin-bottom: 70px !important;
        }

        .modo-terreno-action-card {
            min-height: 64px !important;
            border-radius: 16px !important;
            font-size: 1rem !important;
        }

        /* Evitar que cualquier componente ensanche la pantalla */
        [data-testid="stHorizontalBlock"],
        [data-testid="column"],
        [data-testid="stVerticalBlock"],
        [data-testid="stElementContainer"] {
            min-width: 0 !important;
            max-width: 100% !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Tarjetas compactas para teléfono
st.markdown(
    """
    <style>
    @media screen and (max-width: 600px) {

        /* Menos separación entre las tarjetas */
        div[data-testid="stHorizontalBlock"]:has(
            .modo-terreno-selector-card
        ) {
            gap: 0.45rem !important;
        }

        div[data-testid="stVerticalBlock"]:has(
            .modo-terreno-selector-card
        ) {
            gap: 0.35rem !important;
        }

        /* Tarjeta completa */
        .modo-terreno-selector-card {
            min-height: 0 !important;
            padding: 7px 9px 6px 9px !important;
            margin: 0 0 5px 0 !important;
            border-radius: 14px !important;
        }

        /* Parte superior de la tarjeta */
        .modo-terreno-filter-tools {
            min-height: 0 !important;
            gap: 8px !important;
            margin: 0 !important;
            padding: 0 !important;
        }

        /* Ícono */
        .modo-terreno-selector-icon {
            width: 40px !important;
            height: 40px !important;
            min-width: 40px !important;
            flex-basis: 40px !important;
            border-radius: 11px !important;
        }

        .modo-terreno-selector-emoji {
            font-size: 1.35rem !important;
            line-height: 1 !important;
        }

        /* Título y subtítulo */
        .modo-terreno-selector-title {
            font-size: 0.98rem !important;
            line-height: 1.1 !important;
            margin: 0 !important;
        }

        .modo-terreno-selector-value {
            font-size: 0.74rem !important;
            line-height: 1.15 !important;
            margin-top: 1px !important;
        }

        /* Selector de opciones */
        div[data-baseweb="select"] {
            margin-top: 3px !important;
        }

        div[data-baseweb="select"] > div {
            min-height: 42px !important;
            height: 42px !important;
            border-radius: 10px !important;
            padding-top: 0 !important;
            padding-bottom: 0 !important;
        }

        div[data-baseweb="select"] span {
            font-size: 0.88rem !important;
        }

        /* Menos espacio entre componentes de Streamlit */
        div[data-testid="stElementContainer"]:has(
            .modo-terreno-selector-card
        ) {
            margin-bottom: 0 !important;
        }

        div[data-testid="stVerticalBlockBorderWrapper"]:has(
            .modo-terreno-selector-card
        ) > div {
            padding: 7px 8px !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Compactación extra de tarjetas en teléfono
st.markdown(
    """
    <style>
    @media screen and (max-width: 600px) {

        div[data-testid="stVerticalBlockBorderWrapper"]:has(
            .modo-terreno-selector-card
        ) {
            margin-bottom: 6px !important;
            border-radius: 14px !important;
        }

        div[data-testid="stVerticalBlockBorderWrapper"]:has(
            .modo-terreno-selector-card
        ) > div {
            padding: 5px 7px 7px 7px !important;
        }

        .modo-terreno-selector-card {
            padding: 5px 7px !important;
            margin: 0 0 3px 0 !important;
            border-radius: 12px !important;
        }

        .modo-terreno-filter-tools {
            min-height: 48px !important;
            gap: 7px !important;
        }

        .modo-terreno-selector-icon {
            width: 34px !important;
            height: 34px !important;
            min-width: 34px !important;
            flex-basis: 34px !important;
            border-radius: 9px !important;
        }

        .modo-terreno-selector-emoji {
            font-size: 1.18rem !important;
        }

        .modo-terreno-selector-title {
            font-size: 0.94rem !important;
            line-height: 1.05 !important;
        }

        .modo-terreno-selector-value {
            font-size: 0.69rem !important;
            line-height: 1.05 !important;
        }

        div[data-baseweb="select"] {
            margin-top: 1px !important;
        }

        div[data-baseweb="select"] > div {
            min-height: 38px !important;
            height: 38px !important;
            border-radius: 9px !important;
        }

        div[data-baseweb="select"] span {
            font-size: 0.84rem !important;
        }

        /* Menos espacio vertical entre tarjetas */
        div[data-testid="stVerticalBlock"]:has(
            .modo-terreno-selector-card
        ) {
            gap: 0.18rem !important;
        }

        /* Evitar que botones flotantes tapen el contenido */
        .block-container {
            padding-bottom: 8rem !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Reducir espacio debajo de las pestañas en teléfono
st.markdown(
    """
    <style>
    @media screen and (max-width: 600px) {

        div[data-testid="stTabs"] {
            margin-top: 0 !important;
            margin-bottom: 0 !important;
        }

        div[data-testid="stTabs"] [data-baseweb="tab-list"] {
            margin-bottom: 0 !important;
            padding-bottom: 0 !important;
        }

        div[data-testid="stTabs"] [data-baseweb="tab-panel"] {
            padding-top: 0.35rem !important;
        }

        div[data-testid="stTabs"]
        [data-baseweb="tab-panel"]
        > div {
            padding-top: 0 !important;
            margin-top: 0 !important;
        }

        div[data-testid="stVerticalBlock"]:has(
            .modo-terreno-app-shell-marker
        ) {
            margin-top: 0 !important;
            padding-top: 4px !important;
        }

        div[data-testid="stVerticalBlockBorderWrapper"]:has(
            .modo-terreno-app-shell-marker
        ) {
            margin-top: 0 !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Espacio inferior para evitar que la marca de Streamlit tape contenido
st.markdown(
    """
    <style>
    @media screen and (max-width: 600px) {
        .block-container {
            padding-bottom: 13rem !important;
        }

        div[data-testid="stAppViewBlockContainer"] {
            padding-bottom: 13rem !important;
        }

        .modo-terreno-action-grid {
            margin-bottom: 8rem !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Ajuste inferior para navegación móvil
st.markdown(
    """
    <style>
    @media screen and (max-width: 600px) {

        /* Intentar ocultar elementos inferiores de Streamlit */
        footer,
        [data-testid="stFooter"],
        [data-testid="stBottom"],
        [data-testid="stBottomBlockContainer"],
        [data-testid="stStatusWidget"] {
            display: none !important;
            visibility: hidden !important;
            height: 0 !important;
            min-height: 0 !important;
        }

        /* Espacio seguro para Safari y Chrome móvil */
        .block-container,
        div[data-testid="stAppViewBlockContainer"] {
            padding-bottom: calc(10rem + env(safe-area-inset-bottom)) !important;
        }

        /* Texto bajo el selector de malezas */
        div[data-testid="stCaptionContainer"] {
            margin-top: 0.2rem !important;
            margin-bottom: 0.35rem !important;
        }

        div[data-testid="stCaptionContainer"] p {
            font-size: 0.82rem !important;
            line-height: 1.3 !important;
        }

        /* Compactar el sector del buscador */
        div[data-testid="stTextInput"] {
            margin-top: 0.15rem !important;
            margin-bottom: 0.3rem !important;
        }

        div[data-testid="stTextInput"] input {
            min-height: 42px !important;
            font-size: 0.92rem !important;
        }

        /* Dejar libres los últimos controles */
        div[data-testid="stButton"],
        div[data-testid="stDownloadButton"] {
            scroll-margin-bottom: 9rem !important;
        }

        .modo-terreno-action-grid {
            padding-bottom: 6rem !important;
            margin-bottom: 6rem !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Orden final de datos en las tarjetas de productos móviles
st.markdown(
    """
    <style>
    @media screen and (max-width: 600px) {

        .modo-terreno-product-row-link {
            display: grid !important;
            grid-template-columns: 46px minmax(0, 1fr) !important;
            grid-template-rows: auto auto auto auto !important;
            gap: 7px 10px !important;
            align-items: start !important;
            min-height: 0 !important;
            padding: 12px !important;
            overflow: visible !important;
        }

        /* Ícono */
        .modo-terreno-product-row-link > div:nth-child(1) {
            grid-column: 1 !important;
            grid-row: 1 / 5 !important;
            align-self: start !important;
        }

        /* Nombre e ingrediente */
        .modo-terreno-product-row-link > div:nth-child(2) {
            grid-column: 2 !important;
            grid-row: 1 !important;
            min-width: 0 !important;
        }

        /* Tipo de producto */
        .modo-terreno-product-row-link > div:nth-child(3) {
            grid-column: 2 !important;
            grid-row: 2 !important;
            min-width: 0 !important;
        }

        /* Grupo IRAC / FRAC / HRAC */
        .modo-terreno-product-row-link > div:nth-child(4) {
            grid-column: 2 !important;
            grid-row: 3 !important;
            min-width: 0 !important;
        }

        /* Abejas */
        .modo-terreno-product-row-link > div:nth-child(5) {
            grid-column: 2 !important;
            grid-row: 4 !important;
            min-width: 0 !important;
        }

        .modo-terreno-product-name-box {
            font-size: 1rem !important;
            line-height: 1.2 !important;
            white-space: normal !important;
            overflow-wrap: anywhere !important;
            margin: 0 0 3px 0 !important;
        }

        .modo-terreno-product-name-meta {
            font-size: 0.88rem !important;
            line-height: 1.25 !important;
            white-space: normal !important;
            overflow-wrap: anywhere !important;
        }

        .modo-terreno-pill,
        .modo-terreno-bee {
            display: inline-block !important;
            width: auto !important;
            max-width: 100% !important;
            padding: 5px 10px !important;
            margin: 0 !important;
            border-radius: 999px !important;
            font-size: 0.78rem !important;
            line-height: 1.2 !important;
            white-space: normal !important;
            overflow-wrap: anywhere !important;
        }

        .modo-terreno-lista-meta {
            font-size: 0.82rem !important;
            line-height: 1.3 !important;
            white-space: normal !important;
            overflow-wrap: anywhere !important;
        }

        .modo-terreno-lista-meta strong {
            display: block !important;
            font-size: 0.76rem !important;
            margin-bottom: 2px !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Diseño compacto del cuadro informativo y botones de búsqueda
st.markdown(
    """
    <style>
    .modo-terreno-info-box {
        padding: 6px 12px !important;
        margin: 2px auto 6px auto !important;
        border-width: 1.5px !important;
        border-radius: 13px !important;
        box-shadow: none !important;
    }

    .modo-terreno-info-title {
        font-size: 0.9rem !important;
        line-height: 1.15 !important;
        margin-bottom: 1px !important;
    }

    .modo-terreno-info-subtitle {
        font-size: 0.76rem !important;
        line-height: 1.2 !important;
    }

    .modo-terreno-buscar-marker,
    .modo-terreno-limpiar-marker {
        display: none;
    }

    div[data-testid="stVerticalBlock"]:has(
        .modo-terreno-buscar-marker
    ) div.stButton > button {
        min-height: 46px;
        border: 1px solid #168854 !important;
        border-radius: 13px !important;
        background: linear-gradient(
            135deg,
            #15965b,
            #087944
        ) !important;
        color: white !important;
        font-weight: 850 !important;
        box-shadow: 0 7px 16px rgba(8, 121, 68, 0.20) !important;
    }

    div[data-testid="stVerticalBlock"]:has(
        .modo-terreno-buscar-marker
    ) div.stButton > button:hover {
        background: linear-gradient(
            135deg,
            #10834f,
            #066b3b
        ) !important;
        border-color: #066b3b !important;
        transform: translateY(-1px);
    }

    div[data-testid="stVerticalBlock"]:has(
        .modo-terreno-limpiar-marker
    ) div.stButton > button {
        min-height: 46px;
        border: 1px solid #9dcdb5 !important;
        border-radius: 13px !important;
        background: #f4fbf7 !important;
        color: #17663d !important;
        font-weight: 850 !important;
        box-shadow: none !important;
    }

    div[data-testid="stVerticalBlock"]:has(
        .modo-terreno-limpiar-marker
    ) div.stButton > button:hover {
        border-color: #168854 !important;
        background: #e8f7ef !important;
        color: #0d7545 !important;
    }

    @media screen and (max-width: 600px) {
        .modo-terreno-info-box {
            padding: 5px 9px !important;
            margin: 1px 0 5px 0 !important;
            border-radius: 11px !important;
        }

        .modo-terreno-info-title {
            font-size: 0.84rem !important;
        }

        .modo-terreno-info-subtitle {
            font-size: 0.7rem !important;
        }

        div[data-testid="stTextInput"] label p {
            font-size: 0.84rem !important;
            font-weight: 750 !important;
        }

        div[data-testid="stTextInput"] input {
            min-height: 40px !important;
            font-size: 0.88rem !important;
            border-radius: 11px !important;
        }

        div[data-testid="stHorizontalBlock"]:has(
            .modo-terreno-buscar-marker
        ) {
            gap: 0.5rem !important;
        }

        div[data-testid="stVerticalBlock"]:has(
            .modo-terreno-buscar-marker
        ) div.stButton > button,
        div[data-testid="stVerticalBlock"]:has(
            .modo-terreno-limpiar-marker
        ) div.stButton > button {
            min-height: 42px !important;
            padding: 7px 8px !important;
            border-radius: 11px !important;
            font-size: 0.86rem !important;
            margin-bottom: 0 !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)


def obtener_pdf_bytes(referencia_pdf):
    """Obtiene un PDF desde Supabase Storage o desde la carpeta local."""
    import os
    from urllib.parse import unquote, urlparse

    import requests

    if referencia_pdf is None:
        return None, ""

    referencia = str(referencia_pdf).strip()

    if referencia == "" or referencia.lower() == "nan":
        return None, ""

    if referencia.startswith(("http://", "https://")):
        respuesta = requests.get(referencia, timeout=90)
        respuesta.raise_for_status()

        nombre = unquote(
            urlparse(referencia).path.rsplit("/", 1)[-1]
        )

        return respuesta.content, nombre

    ruta_local = os.path.join("pdfs", referencia)

    if os.path.exists(ruta_local):
        with open(ruta_local, "rb") as archivo:
            return archivo.read(), referencia

    return None, referencia


def mostrar_pdf_bytes(pdf_bytes, alto=650):
    pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")

    visor_pdf = f"""
        <iframe
            src="data:application/pdf;base64,{pdf_base64}"
            width="100%"
            height="{alto}"
            type="application/pdf">
        </iframe>
    """

    st.markdown(visor_pdf, unsafe_allow_html=True)
def generar_ficha_word(fila):
    from datetime import datetime
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output = BytesIO()

    doc = Document()

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    titulo = doc.add_heading("Ficha técnica de producto agrícola", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha = datetime.now().strftime("%d-%m-%Y")
    parrafo_fecha = doc.add_paragraph(f"Fecha de generación: {fecha}")
    parrafo_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    def limpiar(valor):
        valor = str(valor).strip()

        if valor.lower() == "nan":
            return ""

        return valor

    producto = limpiar(fila.get("nombre", ""))
    tipo = limpiar(fila.get("tipo", ""))
    ingrediente = limpiar(fila.get("ingrediente", ""))
    grupo = limpiar(fila.get("grupo", ""))
    dosis = limpiar(fila.get("dosis", ""))
    cultivos = limpiar(fila.get("cultivos", ""))

    enfermedades = limpiar(fila.get("enfermedades", ""))
    insectos = limpiar(fila.get("insectos", ""))

    problema = limpiar(fila.get("problema", ""))

    if problema == "":
        problema = ", ".join(
            dato for dato in [enfermedades, insectos]
            if dato != ""
        )

    carencia = limpiar(fila.get("carencia", ""))
    reingreso = limpiar(fila.get("reingreso", ""))
    abejas = limpiar(fila.get("toxicidad_abejas", ""))
    compatibilidad = limpiar(fila.get("compatibilidad", ""))
    incompatibilidad = limpiar(fila.get("incompatibilidad", ""))
    fitotoxicidad = limpiar(fila.get("fitotoxicidad", ""))
    pdf = limpiar(fila.get("pdf", ""))

    doc.add_heading(producto, level=2)

    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"

    tabla.rows[0].cells[0].text = "Campo"
    tabla.rows[0].cells[1].text = "Información"

    campos = [
        ("Producto", producto),
        ("Tipo", tipo),
        ("Ingrediente activo", ingrediente),
        ("Grupo químico / modo de acción", grupo),
        ("Dosis", dosis),
        ("Cultivos", cultivos),
        ("Problema / plaga / enfermedad", problema),
        ("Compatibilidad", compatibilidad),
        ("Incompatibilidad", incompatibilidad),
        ("Fitotoxicidad", fitotoxicidad),
        ("Carencia", carencia),
        ("Reingreso", reingreso),
        ("Toxicidad para abejas", abejas),
        ("PDF de respaldo", pdf)
    ]

    for campo, valor in campos:
        celdas = tabla.add_row().cells
        celdas[0].text = campo
        celdas[1].text = valor

    doc.add_paragraph("")

    advertencia = doc.add_paragraph()
    advertencia.add_run("Advertencia importante: ").bold = True
    advertencia.add_run(
        "Esta ficha fue generada automáticamente desde IA Agrícola V15. "
        "Antes de aplicar cualquier producto, verificar siempre la etiqueta SAG oficial, "
        "la dosis, el cultivo autorizado, el período de carencia, el reingreso y las condiciones de uso."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Generado por IA Agrícola V15.")

    doc.save(output)

    return output.getvalue()


def nombre_archivo_ficha(nombre):
    nombre = str(nombre).strip().lower()
    nombre = nombre.replace(" ", "_")
    nombre = nombre.replace("/", "_")
    nombre = nombre.replace("\\", "_")
    nombre = nombre.replace(".", "_")

    if nombre == "":
        nombre = "producto"

    return f"ficha_tecnica_{nombre}.docx"


def mostrar_ficha_visual(fila):
    def limpiar(valor):
        valor = str(valor).strip()

        if valor.lower() == "nan":
            return ""

        return valor

    producto = limpiar(fila.get("nombre", ""))
    tipo = limpiar(fila.get("tipo", ""))
    ingrediente = limpiar(fila.get("ingrediente", ""))
    grupo = limpiar(fila.get("grupo", ""))
    dosis = limpiar(fila.get("dosis", ""))
    cultivos = limpiar(fila.get("cultivos", ""))

    problema = limpiar(fila.get("problema", ""))

    if problema == "":
        enfermedades = limpiar(fila.get("enfermedades", ""))
        insectos = limpiar(fila.get("insectos", ""))

        problema = ", ".join(
            dato for dato in [enfermedades, insectos]
            if dato != ""
        )

    carencia = limpiar(fila.get("carencia", ""))
    reingreso = limpiar(fila.get("reingreso", ""))
    abejas = limpiar(fila.get("toxicidad_abejas", ""))
    compatibilidad = limpiar(fila.get("compatibilidad", ""))
    incompatibilidad = limpiar(fila.get("incompatibilidad", ""))
    fitotoxicidad = limpiar(fila.get("fitotoxicidad", ""))
    pdf = limpiar(fila.get("pdf", ""))

    color_tipo = "#9fe8c8"

    if tipo.lower() == "insecticida":
        color_tipo = "#9fe8c8"

    if tipo.lower() == "fungicida":
        color_tipo = "#b8d7ff"

    if tipo.lower() == "herbicida":
        color_tipo = "#d8c2ff"

    if tipo.lower() == "biológico":
        color_tipo = "#c8f7b8"

    html = f"""
    <div style="
        border: 4px solid #9fe8c8;
        border-radius: 40px;
        padding: 28px;
        margin-top: 20px;
        margin-bottom: 25px;
        background: white;
        position: relative;
        font-family: Arial, sans-serif;
    ">

        <div style="
            position: absolute;
            right: 0;
            top: 0;
            bottom: 0;
            width: 70px;
            background: {color_tipo};
            border-radius: 0 35px 35px 0;
            display: flex;
            align-items: center;
            justify-content: center;
        ">
            <div style="
                writing-mode: vertical-rl;
                transform: rotate(180deg);
                color: white;
                font-weight: bold;
                font-size: 24px;
                letter-spacing: 2px;
            ">
                {tipo.upper()}
            </div>
        </div>

        <div style="padding-right: 85px;">
            <h1 style="
                text-align: center;
                font-size: 46px;
                margin: 0;
                color: #111;
                font-family: Georgia, serif;
            ">
                {producto}
            </h1>

            <hr style="border: 1px solid #9fe8c8; margin: 18px 0 24px 0;">

            <div style="
                display: grid;
                grid-template-columns: 1fr 1fr 1fr;
                gap: 22px;
                font-size: 17px;
            ">

                <div>
                    <h3>🧪 Grupo químico</h3>
                    <p>{grupo}</p>

                    <h3>🧬 Composición</h3>
                    <p><b>Ingrediente activo:</b><br>{ingrediente}</p>

                    <h3>🐛 Problema</h3>
                    <p>{problema}</p>
                </div>

                <div>
                    <h3>👩‍🌾 Aplicación</h3>
                    <p><b>Dosis:</b><br>{dosis}</p>

                    <p><b>Cultivos:</b><br>{cultivos}</p>

                    <p><b>Carencia:</b><br>{carencia}</p>
                </div>

                <div>
                    <h3>🔗 Compatibilidad</h3>
                    <p>Sin información cargada todavía.</p>

                    <h3>🚶 Reingreso</h3>
                    <p>{reingreso}</p>

                    <h3>🐝 Abejas</h3>
                    <p style="color:red; font-weight:bold;">{abejas}</p>
                </div>

            </div>

            <div style="
                margin-top: 26px;
                padding: 14px;
                background: #fff8a8;
                border: 2px solid #111;
                width: fit-content;
                font-weight: bold;
                font-size: 20px;
            ">
                Revisar etiqueta SAG oficial antes de aplicar
            </div>

            <p style="margin-top: 18px; font-size: 14px;">
                PDF de respaldo: {pdf}
            </p>
        </div>
    </div>
    """

    st.html(textwrap.dedent(html))




def guardar_compatibilidad_manual(
    id_producto,
    compatibilidad,
    incompatibilidad,
    fitotoxicidad
):
    actualizar_compatibilidad_producto(
        id_producto=id_producto,
        compatibilidad=compatibilidad,
        incompatibilidad=incompatibilidad,
        fitotoxicidad=fitotoxicidad
    )


def convertir_excel(df):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Productos")

    return output.getvalue()



def generar_respaldo_completo(df):
    output = BytesIO()

    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zipf:
        if os.path.exists("agro.db"):
            zipf.write("agro.db", "agro.db")

        if not df.empty:
            excel_respaldo = convertir_excel(
                df.drop(columns=["id"], errors="ignore")
            )

            zipf.writestr(
                "base_productos_agricolas.xlsx",
                excel_respaldo
            )

        if os.path.exists("pdfs"):
            for carpeta, _, archivos in os.walk("pdfs"):
                for archivo in archivos:
                    ruta_archivo = os.path.join(carpeta, archivo)
                    nombre_zip = os.path.relpath(ruta_archivo, ".")

                    zipf.write(
                        ruta_archivo,
                        nombre_zip
                    )

    return output.getvalue()


def completar_tipo_fila(fila):
    tipo_actual = str(fila.get("tipo", "")).strip()

    if tipo_actual != "" and tipo_actual.lower() != "nan":
        return tipo_actual

    grupo = str(fila.get("grupo", "")).lower()
    ingrediente = str(fila.get("ingrediente", "")).lower()
    nombre = str(fila.get("nombre", "")).lower()

    texto = grupo + " " + ingrediente + " " + nombre

    if "biologico" in texto or "biológico" in texto or "bacillus" in texto:
        return "Biológico"

    if "irac" in texto or "pirimicarb" in texto or "imidacloprid" in texto or "profenofos" in texto:
        return "Insecticida"

    if "hrac" in texto or "oxifluorfen" in texto or "penoxsulam" in texto or "simazina":
        return "Herbicida"

    if "frac" in texto or "boscalid" in texto or "fenamidona" in texto or "propamocarb" in texto:
        return "Fungicida"

    if "fungicida" in texto:
        return "Fungicida"

    if "insecticida" in texto:
        return "Insecticida"

    if "herbicida" in texto:
        return "Herbicida"

    return ""


def limpiar_valor(valor):
    valor = str(valor).strip()

    if valor.lower() == "nan":
        return ""

    return valor


def valor_o_sin_info(valor):
    valor = limpiar_valor(valor)

    if valor == "":
        return "Sin información en etiqueta"

    return valor


def escapar_html(valor):
    return html.escape(str(valor_o_sin_info(valor)))


def slug_archivo_producto(valor):
    texto = limpiar_valor(valor).lower()
    texto = re.sub(r"\s+etiqueta\b.*$", "", texto, flags=re.IGNORECASE | re.DOTALL)
    texto = re.sub(r"[^a-z0-9]+", "_", texto)
    return texto.strip("_")


def imagen_producto_html(nombre_producto):
    slug = slug_archivo_producto(nombre_producto)

    if slug == "":
        return '<div class="modo-terreno-product-bottle"></div>'

    carpetas = ["imagenes_productos", "imagenes", "images", "img"]
    extensiones = [".png", ".jpg", ".jpeg", ".webp"]

    for carpeta in carpetas:
        for extension in extensiones:
            ruta = os.path.join(carpeta, slug + extension)

            if os.path.exists(ruta):
                mime = "image/jpeg" if extension in [".jpg", ".jpeg"] else f"image/{extension.lstrip('.')}"

                with open(ruta, "rb") as archivo_imagen:
                    imagen_base64 = base64.b64encode(archivo_imagen.read()).decode("utf-8")

                return (
                    f'<img class="modo-terreno-product-image" '
                    f'src="data:{mime};base64,{imagen_base64}" alt="Producto">'
                )

    return '<div class="modo-terreno-product-bottle"></div>'


def separar_opciones(valor):
    texto = limpiar_valor(valor)

    if texto == "":
        return []

    for separador in ["\n", ";", "|"]:
        texto = texto.replace(separador, ",")

    opciones = []

    for parte in texto.split(","):
        parte = parte.strip()

        if parte != "":
            opciones.append(parte)

    return opciones


def opciones_desde_columna(df, columna):
    if columna not in df.columns:
        return ["Todos"]

    opciones = []

    for valor in df[columna].dropna().tolist():
        opciones.extend(separar_opciones(valor))

    return ["Todos"] + sorted(set(opciones), key=lambda item: item.lower())


def columna_malezas_disponible(df):
    for columna in ["malezas", "maleza", "control_malezas"]:
        if columna in df.columns:
            return columna

    return None


def problema_controlado(fila, columna_malezas=None):
    datos = [
        limpiar_valor(fila.get("enfermedades", "")),
        limpiar_valor(fila.get("insectos", ""))
    ]

    if columna_malezas:
        datos.append(limpiar_valor(fila.get(columna_malezas, "")))

    datos = [dato for dato in datos if dato != ""]

    return ", ".join(datos)


def normalizar_texto(valor):
    valor = limpiar_valor(valor).lower()

    reemplazos = {
        "á": "a",
        "é": "e",
        "í": "i",
        "ó": "o",
        "ú": "u",
        "ñ": "n",
        "®": "",
        "™": "",
        "%": " ",
        ".": " ",
        ",": " ",
        "-": " ",
        "_": " ",
        "/": " "
    }

    for viejo, nuevo in reemplazos.items():
        valor = valor.replace(viejo, nuevo)

    return " ".join(valor.split())


def palabras_producto(nombre):
    palabras_excluir = {
        "sc", "ec", "wg", "wp", "sl", "od", "sp", "sg",
        "etiqueta", "pdf", "de", "del", "la", "el", "los",
        "las", "producto"
    }

    palabras = []

    for palabra in normalizar_texto(nombre).split():
        if palabra in palabras_excluir:
            continue

        if len(palabra) <= 2:
            continue

        palabras.append(palabra)

    return palabras


def texto_menciona_producto(texto, nombre_producto):
    texto = normalizar_texto(texto)
    nombre_producto = normalizar_texto(nombre_producto)

    if texto == "" or nombre_producto == "":
        return False

    if nombre_producto in texto:
        return True

    palabras = palabras_producto(nombre_producto)

    if not palabras:
        return False

    coincidencias = sum(1 for palabra in palabras if palabra in texto)

    if len(palabras) == 1 and coincidencias == 1:
        return True

    if len(palabras) >= 2 and coincidencias >= 2:
        return True

    return False


def filtrar_modo_terreno(
    df,
    cultivo,
    plaga,
    enfermedad,
    maleza,
    busqueda,
    columna_malezas
):
    df_filtrado = df.copy()

    if cultivo != "Todos" and "cultivos" in df_filtrado.columns:
        df_filtrado = df_filtrado[
            df_filtrado["cultivos"].astype(str).str.contains(
                cultivo,
                case=False,
                na=False
            )
        ]

    if plaga != "Todos" and "insectos" in df_filtrado.columns:
        df_filtrado = df_filtrado[
            df_filtrado["insectos"].astype(str).str.contains(
                plaga,
                case=False,
                na=False
            )
        ]

    if enfermedad != "Todos" and "enfermedades" in df_filtrado.columns:
        df_filtrado = df_filtrado[
            df_filtrado["enfermedades"].astype(str).str.contains(
                enfermedad,
                case=False,
                na=False
            )
        ]

    if (
        maleza != "Todos"
        and columna_malezas
        and columna_malezas in df_filtrado.columns
    ):
        df_filtrado = df_filtrado[
            df_filtrado[columna_malezas].astype(str).str.contains(
                maleza,
                case=False,
                na=False
            )
        ]

    if busqueda:
        filtro_busqueda = df_filtrado.apply(
            lambda fila: fila.astype(str).str.contains(
                busqueda,
                case=False,
                na=False
            ).any(),
            axis=1
        )

        df_filtrado = df_filtrado[filtro_busqueda]

    return df_filtrado



def emoji_tipo_producto(tipo):
    texto = limpiar_valor(tipo).lower()

    if "herbicida" in texto:
        return "🌿"

    if "fungicida" in texto and "insecticida" in texto:
        return "🌱"

    if "fungicida" in texto:
        return "🍄"

    if "insecticida" in texto:
        return "🐛"

    if "acaricida" in texto:
        return "🕷️"

    if "bactericida" in texto:
        return "🦠"

    return "🧪"



def tarjeta_producto_modo_terreno(fila, columna_malezas=None):
    nombre = valor_o_sin_info(fila.get("nombre", ""))
    tipo = valor_o_sin_info(fila.get("tipo", ""))
    ingrediente = valor_o_sin_info(fila.get("ingrediente", ""))
    grupo = valor_o_sin_info(fila.get("grupo", ""))
    dosis = valor_o_sin_info(fila.get("dosis", ""))
    cultivos = valor_o_sin_info(fila.get("cultivos", ""))
    problema = valor_o_sin_info(problema_controlado(fila, columna_malezas))
    carencia = valor_o_sin_info(fila.get("carencia", ""))
    reingreso = valor_o_sin_info(fila.get("reingreso", ""))
    abejas = valor_o_sin_info(fila.get("toxicidad_abejas", ""))
    compatibilidad = valor_o_sin_info(fila.get("compatibilidad", ""))
    incompatibilidad = valor_o_sin_info(fila.get("incompatibilidad", ""))
    fitotoxicidad = valor_o_sin_info(fila.get("fitotoxicidad", ""))
    pdf = valor_o_sin_info(fila.get("pdf", ""))

    html = f"""
    <div class="modo-terreno-card">
        <h3>{nombre}</h3>
        <span class="modo-terreno-pill">{tipo}</span>
        <span class="modo-terreno-pill">{grupo}</span>
        <div class="modo-terreno-grid">
            <div>
                <div class="modo-terreno-label">Ingrediente activo</div>
                <div class="modo-terreno-value">{ingrediente}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Dosis</div>
                <div class="modo-terreno-value">{dosis}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Cultivo</div>
                <div class="modo-terreno-value">{cultivos}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Problema que controla</div>
                <div class="modo-terreno-value">{problema}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Carencia</div>
                <div class="modo-terreno-value">{carencia}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Reingreso</div>
                <div class="modo-terreno-value">{reingreso}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Toxicidad para abejas</div>
                <div class="modo-terreno-value">{abejas}</div>
            </div>
            <div>
                <div class="modo-terreno-label">PDF de respaldo</div>
                <div class="modo-terreno-value">{pdf}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Compatibilidad</div>
                <div class="modo-terreno-value">{compatibilidad}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Incompatibilidad</div>
                <div class="modo-terreno-value">{incompatibilidad}</div>
            </div>
            <div>
                <div class="modo-terreno-label">Fitotoxicidad</div>
                <div class="modo-terreno-value">{fitotoxicidad}</div>
            </div>
        </div>
    </div>
    """

    st.markdown(textwrap.dedent(html), unsafe_allow_html=True)


def icono_selector_modo_terreno(titulo):
    iconos = {
        "Cultivo": """
            <svg viewBox="0 0 24 24" aria-hidden="true">
                <path d="M11 20A7 7 0 0 1 4 13c0-5 5-9 15-9 0 10-4 15-9 15Z"></path>
                <path d="M4 20c4-6 8-9 15-12"></path>
            </svg>
        """,
        "Plaga": """
            <svg viewBox="0 0 24 24" aria-hidden="true">
                <path d="M8 2l1.88 1.88"></path>
                <path d="M14.12 3.88 16 2"></path>
                <path d="M9 7.13v-1a3 3 0 0 1 6 0v1"></path>
                <path d="M12 20c-3.3 0-6-2.7-6-6v-3a6 6 0 0 1 12 0v3c0 3.3-2.7 6-6 6Z"></path>
                <path d="M12 20v-9"></path>
                <path d="M6 13H2"></path>
                <path d="M22 13h-4"></path>
                <path d="M6.7 17 3 19"></path>
                <path d="M17.3 17 21 19"></path>
                <path d="M6.7 9 3 7"></path>
                <path d="M17.3 9 21 7"></path>
            </svg>
        """,
        "Enfermedad": """
            <svg viewBox="0 0 24 24" aria-hidden="true">
                <path d="M20 13c0 5-3.5 7.5-8 9-4.5-1.5-8-4-8-9V5l8-3 8 3v8Z"></path>
                <path d="m9 12 2 2 4-5"></path>
            </svg>
        """,
        "Maleza": """
            <svg viewBox="0 0 24 24" aria-hidden="true">
                <path d="M12 22V9"></path>
                <path d="M12 13c-3.5 0-6-2.5-6-6 3.5 0 6 2.5 6 6Z"></path>
                <path d="M12 15c3.5 0 6-2.5 6-6-3.5 0-6 2.5-6 6Z"></path>
                <path d="M5 22h14"></path>
            </svg>
        """
    }

    return textwrap.dedent(iconos.get(titulo, iconos["Cultivo"]))


def tarjeta_selector_modo_terreno(titulo, icono):
    clase_icono = {
        "Cultivo": "selector-cultivo",
        "Plaga": "selector-plaga",
        "Enfermedad": "selector-enfermedad",
        "Maleza": "selector-maleza"
    }.get(titulo, "selector-cultivo")

    html_selector = f"""
    <div class="modo-terreno-selector-card {clase_icono}">
        <div class="modo-terreno-selector-icon">
            <span class="modo-terreno-selector-emoji">{html.escape(icono)}</span>
        </div>
        <div>
            <div class="modo-terreno-selector-title">{html.escape(titulo)}</div>
            <div class="modo-terreno-selector-value">Selecciona una opción</div>
        </div>
    </div>
    """

    st.html(textwrap.dedent(html_selector))


def clase_tipo_producto(tipo):
    tipo_limpio = limpiar_valor(tipo).lower()

    if "fung" in tipo_limpio:
        return "fungicida"

    if "herb" in tipo_limpio:
        return "herbicida"

    if "insect" in tipo_limpio:
        return "insecticida"

    return "insecticida"


def base_datos_mockup_header():
    html_base = """
    <div class="base-app-shell">
        <div class="base-hero">
            <div class="base-hero-icon"></div>
            <div>
                <h2>Base de datos</h2>
                <p>Consulta la información técnica de los productos disponibles.</p>
            </div>
        </div>
    </div>
    """

    st.html(textwrap.dedent(html_base))



# Ajuste de ancho de columnas en Base de datos
st.markdown(
    """
    <style>
    .base-table-head,
    .base-table-row {
        grid-template-columns:
            minmax(220px, 1.45fr)
            minmax(140px, 0.82fr)
            minmax(105px, 0.58fr)
            minmax(220px, 1.15fr)
            !important;
    }

    .base-table-head .base-table-cell:nth-child(2),
    .base-table-row .base-table-cell:nth-child(2),
    .base-table-head .base-table-cell:nth-child(3),
    .base-table-row .base-table-cell:nth-child(3) {
        padding-left: 10px !important;
        padding-right: 10px !important;
    }

    @media screen and (max-width: 700px) {
        .base-table-head,
        .base-table-row {
            grid-template-columns:
                minmax(180px, 1.35fr)
                minmax(125px, 0.78fr)
                minmax(95px, 0.55fr)
                minmax(190px, 1.1fr)
                !important;
            min-width: 650px;
        }

        .base-table-scroll-body {
            overflow-x: auto !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Iconos bonitos por tipo de producto
st.markdown(
    """
    <style>
    .modo-terreno-product-icon,
    .base-product-icon {
        position: relative !important;
        display: inline-flex !important;
        align-items: center !important;
        justify-content: center !important;
        overflow: visible !important;
    }

    .modo-terreno-product-icon::before,
    .base-product-icon::before {
        position: absolute;
        font-size: 1.35rem;
        line-height: 1;
    }

    .modo-terreno-product-icon.herbicida::before,
    .base-product-icon.herbicida::before {
        content: "🌿";
    }

    .modo-terreno-product-icon.fungicida::before,
    .base-product-icon.fungicida::before {
        content: "🍄";
    }

    .modo-terreno-product-icon.insecticida::before,
    .base-product-icon.insecticida::before {
        content: "🐛";
    }

    .modo-terreno-product-icon.acaricida::before,
    .base-product-icon.acaricida::before {
        content: "🕷️";
    }

    .modo-terreno-product-icon.bactericida::before,
    .base-product-icon.bactericida::before {
        content: "🦠";
    }

    .modo-terreno-product-icon.fungicida-insecticida::before,
    .base-product-icon.fungicida-insecticida::before {
        content: "🌱";
    }

    .modo-terreno-product-icon.otro::before,
    .base-product-icon.otro::before {
        content: "🧪";
    }

    @media screen and (max-width: 600px) {
        .modo-terreno-product-icon::before,
        .base-product-icon::before {
            font-size: 1.25rem !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Ajustes móviles simples
st.markdown(
    """
    <style>
    @media screen and (max-width: 700px) {

        /* Evita que cualquier bloque se salga de la pantalla */
        html, body, [data-testid="stAppViewContainer"] {
            overflow-x: hidden !important;
        }

        .block-container {
            padding-left: 0.65rem !important;
            padding-right: 0.65rem !important;
            max-width: 100% !important;
        }

        /* Botones de Streamlit en teléfono */
        .stButton > button,
        .stDownloadButton > button,
        a[data-testid="stBaseButton-secondary"],
        a[data-testid="stLinkButton"] {
            width: 100% !important;
            min-height: 42px !important;
            font-size: 0.86rem !important;
            border-radius: 12px !important;
        }

        /* Selectores más cómodos */
        div[data-baseweb="select"] {
            width: 100% !important;
        }

        /* Expander más compacto */
        details {
            border-radius: 12px !important;
        }

        details summary {
            font-size: 0.9rem !important;
            font-weight: 800 !important;
        }

        /* Ficha técnica compacta móvil */
        .ficha-compacta {
            padding: 10px !important;
            border-radius: 16px !important;
        }

        .ficha-producto {
            gap: 8px !important;
        }

        .ficha-titulo {
            font-size: 1.12rem !important;
            line-height: 1.12 !important;
        }

        .ficha-subtitulo {
            font-size: 0.78rem !important;
        }

        .ficha-abejas-badge {
            font-size: 0.72rem !important;
            padding: 7px 9px !important;
        }

        .ficha-resumen {
            grid-template-columns: 1fr !important;
            width: 100% !important;
        }

        .ficha-item {
            padding: 9px 10px !important;
            border-right: 0 !important;
            border-bottom: 1px solid #e4eaf0 !important;
        }

        .ficha-item:last-child {
            border-bottom: 0 !important;
        }

        .ficha-item-titulo {
            font-size: 0.78rem !important;
        }

        .ficha-item-valor {
            font-size: 0.76rem !important;
        }

        .ficha-detalles {
            grid-template-columns: 1fr !important;
        }

        .ficha-footer {
            font-size: 0.72rem !important;
            line-height: 1.3 !important;
        }

        /* Base de datos: evita que se rompa en celular */
        .base-table-scroll,
        .base-table-wrapper,
        .base-table-container {
            overflow-x: auto !important;
            max-width: 100% !important;
        }

        .base-table-head,
        .base-table-row {
            min-width: 640px !important;
        }

        /* Tarjetas de productos modo terreno */
        .modo-terreno-product-row-link {
            grid-template-columns: 36px 1fr !important;
            gap: 8px !important;
            padding: 9px !important;
        }

        .modo-terreno-product-row-link > div:nth-child(3),
        .modo-terreno-product-row-link > div:nth-child(4),
        .modo-terreno-product-row-link > div:nth-child(5) {
            grid-column: 2 / 3 !important;
        }

        .modo-terreno-product-name-box {
            font-size: 0.9rem !important;
            line-height: 1.15 !important;
        }

        .modo-terreno-product-name-meta {
            font-size: 0.74rem !important;
        }

        .modo-terreno-product-icon {
            width: 32px !important;
            height: 32px !important;
        }

        /* PDF dentro de la app más bajo en teléfono */
        iframe {
            max-width: 100% !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)



# Tarjetas móviles más simples y compactas
st.markdown(
    """
    <style>
    @media screen and (max-width: 700px) {

        .modo-terreno-product-row-link {
            display: grid !important;
            grid-template-columns: 46px 1fr auto !important;
            gap: 8px 10px !important;
            align-items: center !important;
            padding: 11px !important;
            border-radius: 18px !important;
            min-height: 0 !important;
        }

        .modo-terreno-product-icon-cell {
            grid-column: 1 / 2 !important;
            grid-row: 1 / 4 !important;
            align-self: start !important;
        }

        .modo-terreno-product-icon {
            width: 42px !important;
            height: 42px !important;
            border-radius: 14px !important;
            border: 0 !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            font-size: 1.45rem !important;
            background: #eef8f2 !important;
            box-shadow: none !important;
        }

        .modo-terreno-product-icon::before,
        .modo-terreno-product-icon::after {
            display: none !important;
            content: none !important;
        }

        .modo-terreno-product-row-link > div:nth-child(2) {
            grid-column: 2 / 4 !important;
            grid-row: 1 / 2 !important;
            min-width: 0 !important;
        }

        .modo-terreno-product-name-box {
            font-size: 0.98rem !important;
            line-height: 1.15 !important;
            font-weight: 900 !important;
            margin-bottom: 3px !important;
        }

        .modo-terreno-product-name-meta {
            font-size: 0.78rem !important;
            line-height: 1.25 !important;
            display: -webkit-box !important;
            -webkit-line-clamp: 2 !important;
            -webkit-box-orient: vertical !important;
            overflow: hidden !important;
        }

        .modo-terreno-product-row-link > div:nth-child(3) {
            grid-column: 2 / 3 !important;
            grid-row: 2 / 3 !important;
        }

        .modo-terreno-pill {
            display: inline-flex !important;
            width: auto !important;
            font-size: 0.74rem !important;
            padding: 5px 10px !important;
            margin: 2px 0 !important;
        }

        .modo-terreno-product-row-link > div:nth-child(4) {
            grid-column: 2 / 3 !important;
            grid-row: 3 / 4 !important;
            font-size: 0.74rem !important;
            line-height: 1.2 !important;
            min-width: 0 !important;
        }

        .modo-terreno-product-row-link > div:nth-child(5) {
            grid-column: 3 / 4 !important;
            grid-row: 2 / 4 !important;
            align-self: center !important;
            justify-self: end !important;
        }

        .modo-terreno-bee {
            font-size: 0.72rem !important;
            padding: 5px 8px !important;
            white-space: nowrap !important;
        }

        .modo-terreno-lista-meta strong {
            font-size: 0.72rem !important;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)


def tabla_base_datos_mockup(df, filas_visibles=7):
    if df.empty:
        return

    filas_html = []

    for _, fila in df.iterrows():
        nombre = escapar_html(
            fila.get("nombre", "")
        )

        ingrediente = escapar_html(
            valor_o_sin_info(
                fila.get("ingrediente", "")
            )
        )

        grupo = escapar_html(
            valor_o_sin_info(
                fila.get("grupo", "")
            )
        )

        tipo = escapar_html(
            valor_o_sin_info(
                fila.get("tipo", "")
            )
        )

        clase_tipo = clase_tipo_producto(
            fila.get("tipo", "")
        )

        filas_html.append(f"""
        <div class="base-table-row">
            <div class="base-table-cell base-product-cell">
                <span class="base-product-icon {clase_tipo}"></span>
                <span>{nombre}</span>
            </div>
            <div class="base-table-cell">{ingrediente}</div>
            <div class="base-table-cell">{grupo}</div>
            <div class="base-table-cell">
                <span class="base-type-pill {clase_tipo}">
                    {tipo}
                </span>
            </div>
        </div>
        """)

    cantidad = len(df)

    html_tabla = """
    <style>
        .base-table-card-scroll {
            border-radius: 18px;
            overflow: hidden;
            border: 1px solid #dbe3ec;
            background: #ffffff;
            box-shadow: 0 8px 22px rgba(15, 23, 42, 0.06);
        }

        .base-table-scroll-body {
            max-height: 357px;
            overflow-y: auto;
            overflow-x: hidden;
            scrollbar-width: thin;
            scrollbar-color: #94a3b8 #eef3f7;
        }

        .base-table-scroll-body::-webkit-scrollbar {
            width: 8px;
        }

        .base-table-scroll-body::-webkit-scrollbar-track {
            background: #eef3f7;
            border-radius: 10px;
        }

        .base-table-scroll-body::-webkit-scrollbar-thumb {
            background: #94a3b8;
            border-radius: 10px;
        }

        .base-table-scroll-body::-webkit-scrollbar-thumb:hover {
            background: #64748b;
        }

        .base-table-count {
            padding: 6px 14px 8px 14px;
            color: #64748b;
            font-size: 0.76rem;
            text-align: right;
            border-top: 1px solid #edf1f5;
            background: #fafcfd;
        }

        @media screen and (max-width: 700px) {
            .base-table-scroll-body {
                max-height: 350px;
                overflow-x: auto;
            }

            .base-table-card-scroll {
                border-radius: 14px;
            }
        }
    </style>

    <div class="base-app-shell">
        <div class="base-table-card-scroll">
            <div class="base-table-head">
                <div class="base-table-cell">Producto</div>
                <div class="base-table-cell">Ingrediente activo</div>
                <div class="base-table-cell">Grupo</div>
                <div class="base-table-cell">Tipo</div>
            </div>

            <div class="base-table-scroll-body">
                __FILAS__
            </div>

            <div class="base-table-count">
                __CANTIDAD__ producto(s) guardado(s)
            </div>
        </div>
    </div>
    """

    html_tabla = html_tabla.replace(
        "__FILAS__",
        "".join(filas_html)
    )

    html_tabla = html_tabla.replace(
        "__CANTIDAD__",
        str(cantidad)
    )

    st.html(textwrap.dedent(html_tabla))



def acciones_base_datos_mockup():
    html_acciones = """
    <div class="base-app-shell">
        <div class="base-actions-grid">
            <div class="base-select-card">
                <div class="base-product-cell">
                    <span class="base-product-icon"></span>
                    <div>
                        <div style="font-size:1.3rem;font-weight:850;color:#0f172a;">Seleccionar producto</div>
                        <div style="font-size:0.9rem;color:#64748b;margin-top:4px;">Busca y selecciona un producto de la base de datos.</div>
                    </div>
                </div>
                <div class="base-search-box">
                    <span style="font-size:1.35rem;">⌕</span>
                    <span>Respuesta a nombre del producto</span>
                </div>
            </div>
            <div class="base-pdf-stack-card">
                <div class="base-pdf-button">
                    <span class="base-pdf-icon"></span>
                    <span>Ver PDF del<br>producto seleccionado</span>
                </div>
                <div class="base-pdf-button">
                    <span class="base-pdf-icon"></span>
                    <span>Descargar PDF<br>seleccionado</span>
                </div>
            </div>
        </div>
    </div>
    """

    st.html(textwrap.dedent(html_acciones))


def clase_toxicidad_abejas(valor):
    texto = limpiar_valor(valor).lower()

    if "altamente" in texto or "muy" in texto:
        return "alta", "Abejas: alto"

    if texto == "" or texto == "nan":
        return "media", "Abejas: s/i"

    if "no tóxico" in texto or "no toxico" in texto or "virtualmente" in texto or "prácticamente" in texto or "practicamente" in texto:
        return "baja", "Abejas: bajo"

    return "media", "Abejas: revisar"


def badge_advertencia_producto(fila):
    texto = " ".join([
        limpiar_valor(fila.get("toxicidad_abejas", "")),
        limpiar_valor(fila.get("incompatibilidad", "")),
        limpiar_valor(fila.get("fitotoxicidad", ""))
    ]).lower()

    palabras_alerta = [
        "altamente",
        "tóxico",
        "toxico",
        "nocivo",
        "peligroso",
        "incompatible",
        "fitotóxico",
        "fitotoxico"
    ]

    if any(palabra in texto for palabra in palabras_alerta):
        return "NOCIVO"

    return "PRECAUCIÓN"


def texto_corto(valor, largo=95):
    valor = valor_o_sin_info(valor)

    if len(valor) <= largo:
        return valor

    return valor[:largo].rstrip() + "..."


def recuadro_info_modo_terreno():
    html_info = """
    <div class="modo-terreno-info-box">
        <div class="modo-terreno-info-title">Aquí aparecen los productos que se pueden utilizar</div>
        <div class="modo-terreno-info-subtitle">Toca un producto para ver su ficha técnica y recomendaciones.</div>
    </div>
    """

    st.html(textwrap.dedent(html_info))


def lista_compacta_modo_terreno(df, columna_malezas=None, limite=12):
    with st.container(border=True):
        st.markdown(
            '<div class="modo-terreno-lista-botones-marker"></div>',
            unsafe_allow_html=True
        )
        total_productos = min(len(df), limite)
        encabezado_lista = f"""
        <div class="modo-terreno-products-head">
            <div class="modo-terreno-products-title">
                Productos recomendados ({total_productos})
            </div>
            <div class="modo-terreno-products-hint">Toca un producto para ver su ficha</div>
        </div>
        """

        st.markdown(textwrap.dedent(encabezado_lista), unsafe_allow_html=True)

        altura_lista = 320 if len(df) > 4 else None
        contenedor_filas = (
            st.container(height=altura_lista, border=False)
            if altura_lista
            else st.container(border=False)
        )

        with contenedor_filas:
            for _, fila in df.head(limite).iterrows():
                producto_id_raw = str(fila.get("id", ""))
                producto_id = html.escape(producto_id_raw)
                nombre = valor_o_sin_info(fila.get("nombre", ""))
                tipo = valor_o_sin_info(fila.get("tipo", ""))
                ingrediente = texto_corto(fila.get("ingrediente", ""), 70)
                grupo = valor_o_sin_info(fila.get("grupo", ""))
                clase_abejas, texto_abejas = clase_toxicidad_abejas(
                    fila.get("toxicidad_abejas", "")
                )
                clase_tipo = clase_tipo_producto(tipo)

                html_fila = f"""
                <div class="modo-terreno-product-click-row-marker"></div>
                <div class="modo-terreno-product-row-link">
                    <div class="modo-terreno-product-icon-cell">
                        <div class="modo-terreno-product-icon {html.escape(clase_tipo)}">{emoji_tipo_producto(tipo)}</div>
                    </div>
                    <div>
                        <div class="modo-terreno-product-name-box">{html.escape(nombre)}</div>
                        <div class="modo-terreno-product-name-meta">{html.escape(ingrediente)}</div>
                    </div>
                    <div>
                        <span class="modo-terreno-pill">{html.escape(tipo)}</span>
                    </div>
                    <div class="modo-terreno-lista-meta">
                        <strong>IRAC/FRAC/HRAC</strong><br>{html.escape(grupo)}
                    </div>
                    <div>
                        <span class="modo-terreno-bee {clase_abejas}">{html.escape(texto_abejas)}</span>
                    </div>
                </div>
                """

                with st.container(border=False):
                    st.markdown(textwrap.dedent(html_fila), unsafe_allow_html=True)

                    if st.button(
                        f"Ver ficha de {nombre}",
                        key=f"modo_terreno_fila_completa_{producto_id_raw}",
                        use_container_width=True
                    ):
                        st.session_state["modo_terreno_producto_id"] = producto_id_raw
                        st.rerun()

    if len(df) > limite:
        st.caption(
            f"Se muestran {limite} productos en la lista compacta. "
            "Ajusta los filtros para acotar los resultados."
        )

    return st.session_state.get("modo_terreno_producto_id", "")


def lista_compacta_modo_terreno_html_referencia(df, columna_malezas=None, limite=12):
    filas_html = []

    for _, fila in df.head(limite).iterrows():
        producto_id = html.escape(str(fila.get("id", "")))
        nombre = escapar_html(fila.get("nombre", ""))
        tipo = escapar_html(fila.get("tipo", ""))
        ingrediente = escapar_html(texto_corto(fila.get("ingrediente", ""), 70))
        grupo = escapar_html(fila.get("grupo", ""))
        clase_abejas, texto_abejas = clase_toxicidad_abejas(
            fila.get("toxicidad_abejas", "")
        )

        filas_html.append(f"""
        <a class="modo-terreno-lista-item" href="?producto_terreno={producto_id}">
            <div class="modo-terreno-product-icon">✓</div>
            <div>
                <div class="modo-terreno-lista-nombre">{nombre}</div>
                <div class="modo-terreno-lista-meta">{ingrediente}</div>
            </div>
            <div class="modo-terreno-lista-row">
                <div class="modo-terreno-lista-meta">
                    <span class="modo-terreno-pill">{tipo}</span>
                </div>
                <div class="modo-terreno-lista-meta"><strong>IRAC/FRAC/HRAC</strong><br>{grupo}</div>
                <div class="modo-terreno-lista-meta">
                    <span class="modo-terreno-bee {clase_abejas}">{html.escape(texto_abejas)}</span>
                </div>
            </div>
            <div class="modo-terreno-arrow">›</div>
        </a>
        """)

    html_lista = f"""
    <div class="modo-terreno-lista">
        {''.join(filas_html)}
    </div>
    """

    st.html(textwrap.dedent(html_lista))

    if len(df) > limite:
        st.caption(
            f"Se muestran {limite} productos en la lista compacta. "
            "Ajusta los filtros para acotar los resultados."
        )


def bloque_ficha_modo_terreno(titulo, icono, valor, full=False):
    clase_full = " full" if full else ""

    return f"""
    <div class="modo-terreno-ficha-section{clase_full}">
        <div class="modo-terreno-ficha-section-title">
            <span>{html.escape(icono)}</span>
            <span>{html.escape(titulo)}</span>
        </div>
        <div class="modo-terreno-ficha-section-value">{escapar_html(valor)}</div>
    </div>
    """


def ficha_didactica_modo_terreno(
    fila,
    columna_malezas=None
):
    nombre_crudo = limpiar_valor(
        fila.get("nombre", "")
    )

    nombre_visible = re.sub(
        r"\s+ETIQUETA\b.*$",
        "",
        nombre_crudo,
        flags=re.IGNORECASE | re.DOTALL
    ).strip()

    if nombre_visible == "":
        nombre_visible = nombre_crudo

    def valor_ficha(clave, predeterminado="Sin información"):
        valor = limpiar_valor(
            fila.get(clave, "")
        )

        if valor == "" or valor.lower() == "nan":
            return predeterminado

        return valor

    def resumir_texto(valor, maximo=150):
        valor = limpiar_valor(valor)

        if valor == "" or valor.lower() == "nan":
            return "Sin información"

        if len(valor) <= maximo:
            return valor

        corte = valor[:maximo].rsplit(" ", 1)[0]

        if corte == "":
            corte = valor[:maximo]

        return corte.rstrip(" ,.;:") + "…"

    nombre = escapar_html(nombre_visible)
    tipo_texto = valor_ficha("tipo")
    tipo = escapar_html(tipo_texto)
    subtitulo = tipo

    grupo_crudo = valor_ficha("grupo")
    ingrediente_crudo = valor_ficha("ingrediente")
    problema_crudo = problema_controlado(
        fila,
        columna_malezas
    )
    dosis_cruda = valor_ficha("dosis")

    compatibilidad_cruda = valor_ficha(
        "compatibilidad"
    )
    incompatibilidad_cruda = valor_ficha(
        "incompatibilidad"
    )
    reingreso_crudo = valor_ficha(
        "reingreso"
    )
    fitotoxicidad_cruda = valor_ficha(
        "fitotoxicidad"
    )
    toxicidad_abejas_cruda = valor_ficha(
        "toxicidad_abejas"
    )

    grupo = escapar_html(
        resumir_texto(grupo_crudo, 90)
    )

    ingrediente = escapar_html(
        resumir_texto(ingrediente_crudo, 130)
    )

    problema = escapar_html(
        resumir_texto(problema_crudo, 240)
    )

    dosis = escapar_html(
        resumir_texto(dosis_cruda, 110)
    )

    reingreso_resumido = escapar_html(
        resumir_texto(reingreso_crudo, 95)
    )

    compatibilidad = escapar_html(
        compatibilidad_cruda
    )

    incompatibilidad = escapar_html(
        incompatibilidad_cruda
    )

    reingreso = escapar_html(
        reingreso_crudo
    )

    fitotoxicidad = escapar_html(
        fitotoxicidad_cruda
    )

    observaciones_base = limpiar_valor(
        fila.get("observaciones", "")
    )

    observaciones_partes = []

    if observaciones_base:
        observaciones_partes.append(
            observaciones_base
        )

    try:
        from database import obtener_experiencias_producto

        producto_id = limpiar_valor(
            fila.get("id", "")
        )

        if producto_id:
            experiencias = (
                obtener_experiencias_producto(
                    int(producto_id)
                )
            )

            if (
                experiencias is not None
                and not experiencias.empty
            ):
                for _, experiencia in experiencias.head(
                    4
                ).iterrows():
                    partes = []

                    cultivo = limpiar_valor(
                        experiencia.get(
                            "cultivo",
                            ""
                        )
                    )

                    problema_exp = limpiar_valor(
                        experiencia.get(
                            "problema",
                            ""
                        )
                    )

                    dosis_exp = limpiar_valor(
                        experiencia.get(
                            "dosis_usada",
                            ""
                        )
                    )

                    resultado = limpiar_valor(
                        experiencia.get(
                            "resultado_observado",
                            ""
                        )
                    )

                    comentario = limpiar_valor(
                        experiencia.get(
                            "comentario",
                            ""
                        )
                    )

                    if cultivo:
                        partes.append(cultivo)

                    if problema_exp:
                        partes.append(problema_exp)

                    if dosis_exp:
                        partes.append(
                            f"Dosis: {dosis_exp}"
                        )

                    if (
                        resultado
                        and resultado != "Seleccionar"
                    ):
                        partes.append(
                            f"Resultado: {resultado}"
                        )

                    if comentario:
                        partes.append(comentario)

                    if partes:
                        observaciones_partes.append(
                            " • ".join(partes)
                        )

    except Exception:
        pass

    if observaciones_partes:
        observaciones_crudas = "\n".join(
            observaciones_partes
        )
    else:
        observaciones_crudas = (
            "Sin observaciones registradas"
        )

    observaciones_html = escapar_html(
        observaciones_crudas
    ).replace("\\n", "<br>")

    producto_visual = imagen_producto_html(
        nombre_crudo
    )

    advertencia = html.escape(
        badge_advertencia_producto(fila)
    )

    clase_abejas, texto_abejas = (
        clase_toxicidad_abejas(
            toxicidad_abejas_cruda
        )
    )

    texto_abejas_header = texto_abejas

    if clase_abejas == "alta":
        texto_abejas_header = (
            "Tóxico para abejas"
        )

    html_ficha = f"""
    <!doctype html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            * {{
                box-sizing: border-box;
            }}

            body {{
                margin: 0;
                font-family: Arial, sans-serif;
                background: transparent;
                color: #111827;
            }}

            .ficha-compacta {{
                border: 1px solid #d7e2dc;
                border-radius: 22px;
                padding: 14px;
                background: #ffffff;
                box-shadow: 0 8px 22px rgba(15, 23, 42, 0.08);
                width: 100%;
            }}

            .ficha-encabezado {{
                display: flex;
                align-items: flex-start;
                justify-content: space-between;
                gap: 14px;
                padding-bottom: 12px;
                border-bottom: 1px solid #e7edf1;
            }}

            .ficha-producto {{
                display: flex;
                gap: 14px;
                align-items: flex-start;
                min-width: 0;
            }}

            .ficha-titulo {{
                font-size: 1.55rem;
                font-weight: 900;
                color: #0f172a;
                line-height: 1.08;
                overflow-wrap: anywhere;
            }}

            .ficha-subtitulo {{
                margin-top: 4px;
                color: #57606a;
                font-size: 0.9rem;
            }}

            .ficha-abejas {{
                display: flex;
                align-items: center;
                gap: 8px;
                flex: 0 0 auto;
            }}

            .ficha-abejas-badge {{
                border-radius: 10px;
                padding: 9px 13px;
                background: #fff1d8;
                color: #dc531a;
                font-size: 0.82rem;
                font-weight: 900;
                text-transform: uppercase;
            }}

            .ficha-resumen {{
                display: grid;
                grid-template-columns:
                    minmax(150px, 0.8fr)
                    minmax(180px, 1fr)
                    minmax(230px, 1.35fr)
                    minmax(150px, 0.8fr);
                border: 1px solid #e4eaf0;
                border-radius: 15px;
                overflow: hidden;
                margin-top: 12px;
            }}

            .ficha-item {{
                padding: 12px;
                border-right: 1px solid #e4eaf0;
                min-width: 0;
            }}

            .ficha-item:last-child {{
                border-right: 0;
            }}

            .ficha-item-titulo {{
                display: flex;
                align-items: center;
                gap: 7px;
                font-size: 0.82rem;
                font-weight: 850;
                color: #172033;
                margin-bottom: 6px;
            }}

            .ficha-item-valor {{
                font-size: 0.8rem;
                color: #263449;
                line-height: 1.35;
                overflow-wrap: anywhere;
            }}

            .ficha-dose-chip {{
                display: inline-block;
                margin-top: 7px;
                padding: 5px 9px;
                border-radius: 9px;
                background: #e8f7ef;
                color: #17663d;
                font-weight: 800;
                font-size: 0.74rem;
            }}

            .ficha-reingreso {{
                margin-top: 10px;
                border-radius: 12px;
                border: 1px solid #dce9f9;
                background: #f6faff;
                padding: 9px 12px;
                font-size: 0.8rem;
                line-height: 1.35;
            }}

            .ficha-reingreso strong {{
                color: #155fb0;
                margin-right: 6px;
            }}

            .ficha-detalles {{
                margin-top: 10px;
                display: grid;
                grid-template-columns:
                    repeat(2, minmax(0, 1fr));
                gap: 8px;
            }}

            .ficha-detalles details {{
                border: 1px solid #e1e7ed;
                border-radius: 11px;
                background: #fafcfd;
                overflow: hidden;
            }}

            .ficha-detalles summary {{
                cursor: pointer;
                padding: 9px 11px;
                font-size: 0.79rem;
                font-weight: 850;
                color: #263449;
                list-style-position: inside;
            }}

            .ficha-detalle-texto {{
                padding: 0 12px 11px 12px;
                color: #435166;
                font-size: 0.77rem;
                line-height: 1.4;
                overflow-wrap: anywhere;
            }}

            .ficha-footer {{
                margin-top: 10px;
                border-radius: 12px;
                background: #fff8ea;
                border: 1px solid #f3d38c;
                padding: 8px 11px;
                display: flex;
                justify-content: space-between;
                align-items: center;
                gap: 10px;
                color: #79510a;
                font-size: 0.78rem;
                font-weight: 800;
            }}

            @media screen and (max-width: 850px) {{
                .ficha-resumen {{
                    grid-template-columns:
                        repeat(2, minmax(0, 1fr));
                }}

                .ficha-item:nth-child(2) {{
                    border-right: 0;
                }}

                .ficha-item:nth-child(1),
                .ficha-item:nth-child(2) {{
                    border-bottom: 1px solid #e4eaf0;
                }}
            }}

            @media screen and (max-width: 600px) {{
                .ficha-encabezado {{
                    display: block;
                }}

                .ficha-abejas {{
                    margin-top: 10px;
                }}

                .ficha-titulo {{
                    font-size: 1.25rem;
                }}

                .ficha-resumen {{
                    grid-template-columns: 1fr;
                }}

                .ficha-item {{
                    border-right: 0;
                    border-bottom: 1px solid #e4eaf0;
                }}

                .ficha-item:last-child {{
                    border-bottom: 0;
                }}

                .ficha-detalles {{
                    grid-template-columns: 1fr;
                }}

                .ficha-footer {{
                    display: block;
                }}
            }}
        </style>
    </head>

    <body>
        <div class="ficha-compacta">
            <div class="ficha-encabezado">
                <div class="ficha-producto">
                    {producto_visual}

                    <div>
                        <div class="ficha-titulo">
                            {nombre}
                        </div>

                        <div class="ficha-subtitulo">
                            {subtitulo}
                        </div>
                    </div>
                </div>

                <div class="ficha-abejas">
                    <span>🐝</span>

                    <span class="ficha-abejas-badge">
                        {html.escape(texto_abejas_header)}
                    </span>
                </div>
            </div>

            <div class="ficha-resumen">
                <div class="ficha-item">
                    <div class="ficha-item-titulo">
                        ⚗ Grupo químico
                    </div>

                    <div class="ficha-item-valor">
                        {grupo}
                    </div>
                </div>

                <div class="ficha-item">
                    <div class="ficha-item-titulo">
                        🧬 Composición
                    </div>

                    <div class="ficha-item-valor">
                        {ingrediente}
                    </div>
                </div>

                <div class="ficha-item">
                    <div class="ficha-item-titulo">
                        ✹ Plagas, enfermedades o malezas
                    </div>

                    <div class="ficha-item-valor">
                        {problema}
                    </div>
                </div>

                <div class="ficha-item">
                    <div class="ficha-item-titulo">
                        ◖ Aplicación
                    </div>

                    <div class="ficha-item-valor">
                        {dosis}
                    </div>

                    <span class="ficha-dose-chip">
                        Dosis referencial
                    </span>
                </div>
            </div>

            <div class="ficha-reingreso">
                <strong>🚶 Reingreso:</strong>
                {reingreso_resumido}
            </div>

            <div class="ficha-detalles">
                <details>
                    <summary>
                        Compatibilidad
                    </summary>

                    <div class="ficha-detalle-texto">
                        {compatibilidad}
                    </div>
                </details>

                <details>
                    <summary>
                        Incompatibilidad
                    </summary>

                    <div class="ficha-detalle-texto">
                        {incompatibilidad}
                    </div>
                </details>

                <details>
                    <summary>
                        Fitotoxicidad
                    </summary>

                    <div class="ficha-detalle-texto">
                        {fitotoxicidad}
                    </div>
                </details>

                <details>
                    <summary>
                        Reingreso completo
                    </summary>

                    <div class="ficha-detalle-texto">
                        {reingreso}
                    </div>
                </details>

                <details>
                    <summary>
                        Observaciones y experiencia de campo
                    </summary>

                    <div class="ficha-detalle-texto">
                        {observaciones_html}
                    </div>
                </details>
            </div>

            <div class="ficha-footer">
                <span>
                    Lea siempre la etiqueta antes de usar el producto.
                </span>

                <span>
                    {advertencia}
                </span>
            </div>
        </div>
    </body>
    </html>
    """

    st.html(
        textwrap.dedent(html_ficha)
    )



def evaluar_compatibilidad_seleccion(productos):
    resultados = []

    for indice_a in range(len(productos)):
        for indice_b in range(indice_a + 1, len(productos)):
            fila_a = productos[indice_a]
            fila_b = productos[indice_b]

            nombre_a = limpiar_valor(fila_a.get("nombre", ""))
            nombre_b = limpiar_valor(fila_b.get("nombre", ""))
            compat_a = limpiar_valor(fila_a.get("compatibilidad", ""))
            compat_b = limpiar_valor(fila_b.get("compatibilidad", ""))
            incompat_a = limpiar_valor(fila_a.get("incompatibilidad", ""))
            incompat_b = limpiar_valor(fila_b.get("incompatibilidad", ""))

            compatible = (
                texto_menciona_producto(compat_a, nombre_b)
                or texto_menciona_producto(compat_b, nombre_a)
            )

            incompatible = (
                texto_menciona_producto(incompat_a, nombre_b)
                or texto_menciona_producto(incompat_b, nombre_a)
            )

            if incompatible:
                estado = "incompatible"
                mensaje = f"{nombre_a} y {nombre_b}: No compatible según etiqueta."
            elif compatible:
                estado = "compatible"
                mensaje = f"{nombre_a} y {nombre_b}: Compatible según etiqueta."
            else:
                estado = "sin_info"
                mensaje = f"{nombre_a} y {nombre_b}: No hay información suficiente en etiqueta."

            resultados.append((estado, mensaje))

    return resultados


def nombre_corto_producto_label(etiqueta):
    nombre = str(etiqueta).split("|")[0].strip()

    if len(nombre) > 38:
        return nombre[:38].rstrip() + "..."

    return nombre


def compatibilidad_slots_modo_terreno(productos_labels):
    slots = []

    for indice in range(4):
        if indice < len(productos_labels):
            texto_slot = nombre_corto_producto_label(productos_labels[indice])
        else:
            texto_slot = f"Insertar producto {indice + 1}"

        slots.append(f"""
        <div class="modo-terreno-compat-slot">
            {html.escape(texto_slot)}
        </div>
        """)

    html_slots = f"""
    <div class="modo-terreno-compat-wrapper">
        <div class="modo-terreno-compat-title">Compatibilidad</div>
        <div class="modo-terreno-compat-slots">
            {''.join(slots)}
        </div>
    </div>
    """

    st.html(textwrap.dedent(html_slots))


def compatibilidad_resultado_modo_terreno(productos_labels, resultados):
    if len(productos_labels) < 2:
        mensaje_resultado = (
            "No hay información suficiente en etiqueta para confirmar esta mezcla. "
            "Revisar etiqueta oficial antes de aplicar."
        )
    elif not resultados or all(estado == "sin_info" for estado, _ in resultados):
        mensaje_resultado = (
            "No hay información suficiente en etiqueta para confirmar esta mezcla. "
            "Revisar etiqueta oficial antes de aplicar."
        )
    else:
        mensajes = [mensaje for _, mensaje in resultados]
        mensaje_resultado = " ".join(mensajes)

    html_resultado = f"""
        <div class="modo-terreno-compat-result">
            {html.escape(mensaje_resultado)}
        </div>
    """

    st.html(textwrap.dedent(html_resultado))


def botones_finales_modo_terreno():
    html_botones = """
    <div class="modo-terreno-action-grid">
        <a class="modo-terreno-action-card pdf" href="?accion_modo_terreno=cargar_pdf" target="_self">
            <span class="modo-terreno-action-icon pdf"></span>
            <span>Cargar PDF</span>
        </a>
        <a class="modo-terreno-action-card campo" href="?accion_modo_terreno=experiencia" target="_self">
            <span class="modo-terreno-action-icon campo"></span>
            <span>Experiencia de campo</span>
        </a>
        <a class="modo-terreno-action-card db" href="?accion_modo_terreno=base_datos" target="_self">
            <span class="modo-terreno-action-icon db"></span>
            <span>Base de datos</span>
        </a>
    </div>
    """

    st.html(textwrap.dedent(html_botones))




def limpiar_nombre_pdf_subida(nombre):
    nombre = os.path.basename(str(nombre))
    nombre = re.sub(r"[^A-Za-z0-9._-]+", "_", nombre)
    nombre = re.sub(r"_+", "_", nombre)
    return nombre.strip("._") or "etiqueta.pdf"


def formulario_cargar_pdf_modo_terreno():
    st.markdown("### Cargar una etiqueta PDF")
    st.caption(
        "El archivo quedará guardado permanentemente en Supabase."
    )

    archivo = st.file_uploader(
        "Selecciona una etiqueta en formato PDF",
        type=["pdf"],
        accept_multiple_files=False,
        key="modo_terreno_archivo_pdf"
    )

    if archivo is None:
        return

    pdf_bytes = archivo.getvalue()
    nombre_original = archivo.name
    nombre_storage = limpiar_nombre_pdf_subida(nombre_original)

    if not nombre_storage.lower().endswith(".pdf"):
        nombre_storage += ".pdf"

    ruta_temporal = None

    try:
        with tempfile.NamedTemporaryFile(
            suffix=".pdf",
            delete=False
        ) as temporal:
            temporal.write(pdf_bytes)
            ruta_temporal = temporal.name

        with st.spinner("Analizando la etiqueta..."):
            texto = extraer_texto_pdf(ruta_temporal)
            resultado = analizar_texto(
                texto,
                nombre_original
            )
            usos = extraer_usos_pdf(ruta_temporal)

            if usos:
                def unicos_usos(clave):
                    valores = []
                    vistos = set()

                    for uso in usos:
                        valor = str(
                            uso.get(clave, "")
                        ).strip()

                        if not valor:
                            continue

                        identificador = valor.casefold()

                        if identificador not in vistos:
                            vistos.add(identificador)
                            valores.append(valor)

                    return valores

                cultivos_tabla = unicos_usos("cultivo")
                problemas_tabla = unicos_usos("problema")
                dosis_tabla = unicos_usos("dosis")

                palabras_insectos = [
                    "mosca",
                    "trips",
                    "pulgón",
                    "pulgon",
                    "chanchito",
                    "polilla",
                    "ácaro",
                    "acaro",
                    "conchuela",
                    "gusano",
                    "larva",
                    "insecto",
                    "cuncunilla",
                    "copitarsia",
                    "drosophila",
                    "pseudococcus",
                    "frankliniella",
                    "nasonovia"
                ]

                insectos_tabla = []
                enfermedades_tabla = []

                for problema in problemas_tabla:
                    problema_minuscula = (
                        problema.casefold()
                    )

                    es_insecto = any(
                        palabra in problema_minuscula
                        for palabra in palabras_insectos
                    )

                    if es_insecto:
                        insectos_tabla.append(problema)
                    else:
                        enfermedades_tabla.append(
                            problema
                        )

                if cultivos_tabla:
                    resultado["cultivos"] = ", ".join(
                        cultivos_tabla
                    )

                if dosis_tabla:
                    resultado["dosis"] = ", ".join(
                        dosis_tabla
                    )

                if insectos_tabla:
                    resultado["insectos"] = ", ".join(
                        insectos_tabla
                    )

                if enfermedades_tabla:
                    resultado["enfermedades"] = ", ".join(
                        enfermedades_tabla
                    )

    except Exception as error:
        st.error(f"No fue posible analizar el PDF: {error}")
        return

    finally:
        if ruta_temporal and os.path.exists(ruta_temporal):
            os.remove(ruta_temporal)

    st.success(
        f"PDF analizado. Se encontraron {len(usos)} usos o dosis."
    )

    with st.expander("Ver PDF antes de guardar"):
        mostrar_pdf_bytes(pdf_bytes, alto=550)

    st.download_button(
        "Descargar PDF seleccionado",
        data=pdf_bytes,
        file_name=nombre_original,
        mime="application/pdf",
        key="descargar_pdf_nuevo"
    )

    with st.form("formulario_guardar_pdf_supabase"):
        st.markdown("#### Revisa y corrige los datos")

        col1, col2 = st.columns(2)

        with col1:
            nombre = st.text_input(
                "Producto",
                value=resultado.get("producto", "")
            )

            ingrediente = st.text_area(
                "Ingrediente activo",
                value=resultado.get("ingrediente", ""),
                height=100
            )

            grupo = st.text_input(
                "Grupo IRAC / FRAC / HRAC",
                value=resultado.get("grupo", "")
            )

            tipo = st.text_input(
                "Tipo de producto",
                value=resultado.get("tipo", "")
            )

            cultivos = st.text_area(
                "Cultivos",
                value=resultado.get("cultivos", ""),
                height=120
            )

            dosis = st.text_area(
                "Dosis general",
                value=resultado.get("dosis", ""),
                height=100
            )

        with col2:
            enfermedades = st.text_area(
                "Enfermedades",
                value=resultado.get("enfermedades", ""),
                height=100
            )

            insectos = st.text_area(
                "Insectos o plagas",
                value=resultado.get("insectos", ""),
                height=100
            )

            reingreso = st.text_area(
                "Reingreso",
                value=resultado.get("reingreso", ""),
                height=80
            )

            carencia = st.text_area(
                "Carencia",
                value=resultado.get("carencia", ""),
                height=80
            )

            toxicidad_abejas = st.text_area(
                "Toxicidad para abejas",
                value=resultado.get(
                    "toxicidad_abejas",
                    ""
                ),
                height=90
            )

        compatibilidad = st.text_area(
            "Compatibilidad",
            value=resultado.get("compatibilidad", ""),
            height=90
        )

        incompatibilidad = st.text_area(
            "Incompatibilidad",
            value=resultado.get("incompatibilidad", ""),
            height=90
        )

        fitotoxicidad = st.text_area(
            "Fitotoxicidad",
            value=resultado.get("fitotoxicidad", ""),
            height=90
        )

        if usos:
            st.markdown("#### Usos y dosis detectados")

            df_usos = pd.DataFrame(usos)

            columnas_usos = [
                "cultivo",
                "problema",
                "dosis",
                "observaciones",
                "pagina"
            ]

            for columna in columnas_usos:
                if columna not in df_usos.columns:
                    df_usos[columna] = ""

            df_usos_editado = st.data_editor(
                df_usos[columnas_usos],
                use_container_width=True,
                num_rows="dynamic",
                hide_index=True,
                key="editor_usos_pdf"
            )
        else:
            st.warning(
                "No se encontraron tablas de dosis automáticamente. "
                "El producto igualmente se puede guardar."
            )
            df_usos_editado = pd.DataFrame()

        confirmar = st.checkbox(
            "Confirmo que revisé los datos de la etiqueta."
        )

        guardar = st.form_submit_button(
            "Guardar PDF y producto",
            use_container_width=True,
            type="primary"
        )

    if not guardar:
        return

    if not confirmar:
        st.warning(
            "Debes confirmar que revisaste los datos."
        )
        return

    if not str(nombre).strip():
        st.error("El campo Producto es obligatorio.")
        return

    try:
        with st.spinner(
            "Subiendo el PDF y guardando los datos..."
        ):
            url_pdf = subir_pdf_storage(
                nombre_storage,
                pdf_bytes
            )

            producto_id = guardar_producto(
                nombre=nombre,
                ingrediente=ingrediente,
                grupo=grupo,
                tipo=tipo,
                cultivos=cultivos,
                enfermedades=enfermedades,
                insectos=insectos,
                dosis=dosis,
                compatibilidad=compatibilidad,
                incompatibilidad=incompatibilidad,
                fitotoxicidad=fitotoxicidad,
                reingreso=reingreso,
                carencia=carencia,
                toxicidad_abejas=toxicidad_abejas,
                pdf=url_pdf
            )

            usos_guardar = []

            if not df_usos_editado.empty:
                for registro in df_usos_editado.to_dict(
                    orient="records"
                ):
                    usos_guardar.append({
                        "cultivo": registro.get(
                            "cultivo",
                            ""
                        ),
                        "problema": registro.get(
                            "problema",
                            ""
                        ),
                        "dosis": registro.get(
                            "dosis",
                            ""
                        ),
                        "observaciones": registro.get(
                            "observaciones",
                            ""
                        ),
                        "pagina": registro.get(
                            "pagina"
                        ),
                    })

            guardar_usos_producto(
                producto_id,
                url_pdf,
                usos_guardar
            )

        st.success(
            f"{nombre} fue guardado correctamente."
        )

        st.session_state[
            "modo_terreno_busqueda_ejecutada"
        ] = False

    except Exception as error:
        st.error(
            "No fue posible guardar el producto: "
            f"{error}"
        )


def formulario_experiencia_campo_modo_terreno(df_productos):
    st.markdown("### Experiencia de campo / Observaciones")

    if df_productos is None or df_productos.empty:
        st.info("Primero debes tener productos cargados en la base de datos.")
        return

    opciones = []
    mapa_productos = {}

    producto_id_actual = st.session_state.get("modo_terreno_producto_id", "")

    for _, fila in df_productos.iterrows():
        producto_id = str(fila.get("id", "")).strip()
        nombre = str(fila.get("nombre", "")).strip()
        tipo = str(fila.get("tipo", "")).strip()

        if nombre and producto_id:
            etiqueta = f"{nombre} | {tipo} | ID {producto_id}"
            opciones.append(etiqueta)
            mapa_productos[etiqueta] = fila

    if not opciones:
        st.info("No hay productos disponibles para registrar experiencia.")
        return

    indice_defecto = 0
    if producto_id_actual:
        for i, etiqueta in enumerate(opciones):
            if etiqueta.endswith(f"ID {producto_id_actual}"):
                indice_defecto = i
                break

    with st.expander("➕ Registrar nueva experiencia de campo", expanded=False):
        producto_sel = st.selectbox(
            "Producto",
            opciones,
            index=indice_defecto,
            key="experiencia_producto_sel"
        )

        fila_producto = mapa_productos[producto_sel]
        producto_id = int(fila_producto.get("id", 0))
        producto_nombre = str(fila_producto.get("nombre", "")).strip()

        col1, col2 = st.columns(2)

        with col1:
            cultivo = st.text_input(
                "Cultivo",
                placeholder="Ej: pimiento, cerezo, tomate",
                key="experiencia_cultivo"
            )

            problema = st.text_input(
                "Problema observado",
                placeholder="Ej: pulgón, oídio, maleza",
                key="experiencia_problema"
            )

            dosis_usada = st.text_input(
                "Dosis usada",
                placeholder="Ej: 200 cc/ha, 0,5 L/ha",
                key="experiencia_dosis"
            )

            fecha = st.date_input(
                "Fecha de aplicación",
                key="experiencia_fecha"
            )

        with col2:
            horario_aplicacion = st.selectbox(
                "Horario de aplicación",
                ["Seleccionar", "Mañana", "Mediodía", "Tarde", "Noche"],
                key="experiencia_horario"
            )

            condicion_climatica = st.text_input(
                "Condición climática",
                placeholder="Ej: fresco, calor, viento, humedad alta",
                key="experiencia_clima"
            )

            resultado_observado = st.selectbox(
                "Resultado observado",
                ["Seleccionar", "Muy bueno", "Bueno", "Regular", "Malo", "Sin evaluar"],
                key="experiencia_resultado"
            )

        comentario = st.text_area(
            "Comentario / observación",
            placeholder="Ej: buen control a los 3 días, aplicar temprano, no usar con mucho calor...",
            key="experiencia_comentario"
        )

        if st.button("Guardar experiencia de campo", key="btn_guardar_experiencia", use_container_width=True):
            from database import guardar_experiencia_campo

            guardar_experiencia_campo(
                producto_id=producto_id,
                producto_nombre=producto_nombre,
                cultivo=cultivo,
                problema=problema,
                dosis_usada=dosis_usada,
                horario_aplicacion=horario_aplicacion,
                condicion_climatica=condicion_climatica,
                resultado_observado=resultado_observado,
                comentario=comentario,
                fecha=str(fecha)
            )

            st.success("Experiencia de campo guardada correctamente.")
            st.rerun()

    with st.expander("📋 Ver últimas experiencias registradas", expanded=False):
        try:
            from database import obtener_todas_experiencias

            df_exp = obtener_todas_experiencias()

            if df_exp.empty:
                st.caption("Aún no hay observaciones registradas.")
            else:
                for _, exp in df_exp.head(10).iterrows():
                    experiencia_id = int(exp.get("id", 0))

                    with st.container(border=True):
                        st.markdown(f"**{exp.get('producto_nombre', '')}**")
                        st.write(
                            f"**Cultivo:** {exp.get('cultivo', '')} | "
                            f"**Problema:** {exp.get('problema', '')} | "
                            f"**Dosis:** {exp.get('dosis_usada', '')}"
                        )
                        st.write(
                            f"**Horario:** {exp.get('horario_aplicacion', '')} | "
                            f"**Resultado:** {exp.get('resultado_observado', '')} | "
                            f"**Fecha:** {exp.get('fecha', '')}"
                        )
                        st.caption(str(exp.get("comentario", "")))

                        col_editar, col_eliminar = st.columns(2)

                        with col_editar:
                            if st.button(
                                "Editar",
                                key=f"editar_experiencia_{experiencia_id}",
                                use_container_width=True
                            ):
                                st.session_state["experiencia_editando_id"] = experiencia_id
                                st.rerun()

                        with col_eliminar:
                            if st.button(
                                "Eliminar",
                                key=f"eliminar_experiencia_{experiencia_id}",
                                use_container_width=True
                            ):
                                from database import eliminar_experiencia_campo

                                eliminar_experiencia_campo(experiencia_id)
                                st.success("Experiencia eliminada correctamente.")
                                st.rerun()

                    if st.session_state.get("experiencia_editando_id") == experiencia_id:
                        st.markdown("#### Editar experiencia")

                        cultivo_edit = st.text_input(
                            "Cultivo",
                            value=str(exp.get("cultivo", "")),
                            key=f"cultivo_edit_{experiencia_id}"
                        )

                        problema_edit = st.text_input(
                            "Problema",
                            value=str(exp.get("problema", "")),
                            key=f"problema_edit_{experiencia_id}"
                        )

                        dosis_edit = st.text_input(
                            "Dosis usada",
                            value=str(exp.get("dosis_usada", "")),
                            key=f"dosis_edit_{experiencia_id}"
                        )

                        horario_edit = st.selectbox(
                            "Horario de aplicación",
                            ["Seleccionar", "Mañana", "Mediodía", "Tarde", "Noche"],
                            index=(
                                ["Seleccionar", "Mañana", "Mediodía", "Tarde", "Noche"].index(str(exp.get("horario_aplicacion", "")))
                                if str(exp.get("horario_aplicacion", "")) in ["Seleccionar", "Mañana", "Mediodía", "Tarde", "Noche"]
                                else 0
                            ),
                            key=f"horario_edit_{experiencia_id}"
                        )

                        clima_edit = st.text_input(
                            "Condición climática",
                            value=str(exp.get("condicion_climatica", "")),
                            key=f"clima_edit_{experiencia_id}"
                        )

                        resultado_edit = st.selectbox(
                            "Resultado observado",
                            ["Seleccionar", "Muy bueno", "Bueno", "Regular", "Malo", "Sin evaluar"],
                            index=(
                                ["Seleccionar", "Muy bueno", "Bueno", "Regular", "Malo", "Sin evaluar"].index(str(exp.get("resultado_observado", "")))
                                if str(exp.get("resultado_observado", "")) in ["Seleccionar", "Muy bueno", "Bueno", "Regular", "Malo", "Sin evaluar"]
                                else 0
                            ),
                            key=f"resultado_edit_{experiencia_id}"
                        )

                        fecha_edit = st.text_input(
                            "Fecha",
                            value=str(exp.get("fecha", "")),
                            key=f"fecha_edit_{experiencia_id}"
                        )

                        comentario_edit = st.text_area(
                            "Comentario",
                            value=str(exp.get("comentario", "")),
                            key=f"comentario_edit_{experiencia_id}"
                        )

                        col_guardar, col_cancelar = st.columns(2)

                        with col_guardar:
                            if st.button(
                                "Guardar cambios",
                                key=f"guardar_edit_{experiencia_id}",
                                use_container_width=True
                            ):
                                from database import actualizar_experiencia_campo

                                actualizar_experiencia_campo(
                                    experiencia_id=experiencia_id,
                                    cultivo=cultivo_edit,
                                    problema=problema_edit,
                                    dosis_usada=dosis_edit,
                                    horario_aplicacion=horario_edit,
                                    condicion_climatica=clima_edit,
                                    resultado_observado=resultado_edit,
                                    comentario=comentario_edit,
                                    fecha=fecha_edit
                                )

                                st.session_state["experiencia_editando_id"] = None
                                st.success("Experiencia actualizada correctamente.")
                                st.rerun()

                        with col_cancelar:
                            if st.button(
                                "Cancelar",
                                key=f"cancelar_edit_{experiencia_id}",
                                use_container_width=True
                            ):
                                st.session_state["experiencia_editando_id"] = None
                                st.rerun()
        except Exception as error:
            st.warning(f"No se pudieron cargar observaciones: {error}")

def limpiar_filtros_modo_terreno():
    st.session_state["modo_terreno_cultivo"] = "Seleccionar"
    st.session_state["modo_terreno_plaga"] = "Seleccionar"
    st.session_state["modo_terreno_enfermedad"] = "Seleccionar"
    st.session_state["modo_terreno_maleza"] = "Seleccionar"
    st.session_state["modo_terreno_busqueda"] = ""
    st.session_state["modo_terreno_producto_id"] = ""
    st.session_state["modo_terreno_busqueda_ejecutada"] = False
    st.session_state["modo_terreno_criterios_aplicados"] = {}


def buscar_productos_modo_terreno():
    cultivo = st.session_state.get("modo_terreno_cultivo", "Seleccionar")
    plaga = st.session_state.get("modo_terreno_plaga", "Seleccionar")
    enfermedad = st.session_state.get("modo_terreno_enfermedad", "Seleccionar")
    maleza = st.session_state.get("modo_terreno_maleza", "Seleccionar")
    busqueda = st.session_state.get("modo_terreno_busqueda", "")
    hay_criterio = any(
        valor != "Seleccionar"
        for valor in [cultivo, plaga, enfermedad, maleza]
    ) or limpiar_valor(busqueda) != ""

    st.session_state["modo_terreno_busqueda_ejecutada"] = hay_criterio
    st.session_state["modo_terreno_producto_id"] = ""
    st.session_state["modo_terreno_criterios_aplicados"] = {
        "cultivo": cultivo,
        "plaga": plaga,
        "enfermedad": enfermedad,
        "maleza": maleza,
        "busqueda": busqueda
    } if hay_criterio else {}


st.set_page_config(
    page_title="IA Agricola V15",
    layout="wide"
)

aplicar_diseno_responsivo()

st.markdown(
    '<div class="app-top-title">IA Agricola V15</div>',
    unsafe_allow_html=True
)

accion_inicial = st.query_params.get(
    "accion_modo_terreno",
    ""
)

pestana_inicial = (
    "Base de datos"
    if accion_inicial == "base_datos"
    else "Modo terreno"
)

tab_modo_terreno, tab_base = st.tabs(
    [
        "Modo terreno",
        "Base de datos"
    ],
    default=pestana_inicial
)


with tab_modo_terreno:
    with st.container(border=False):
        st.markdown(
            '<div class="modo-terreno-app-shell-marker"></div>',
            unsafe_allow_html=True
        )

        df_terreno = obtener_productos()

        columnas_esperadas = [
            "tipo",
            "dosis",
            "compatibilidad",
            "incompatibilidad",
            "fitotoxicidad"
        ]

        for columna in columnas_esperadas:
            if columna not in df_terreno.columns:
                df_terreno[columna] = ""

        if df_terreno.empty:
            st.info("Aun no hay productos guardados. Carga etiquetas PDF para comenzar.")
        else:
            df_terreno["tipo"] = df_terreno.apply(completar_tipo_fila, axis=1)
            columna_malezas = columna_malezas_disponible(df_terreno)

            cultivos_terreno = opciones_desde_columna(df_terreno, "cultivos")
            plagas_terreno = opciones_desde_columna(df_terreno, "insectos")
            enfermedades_terreno = opciones_desde_columna(df_terreno, "enfermedades")

            if columna_malezas:
                malezas_terreno = opciones_desde_columna(df_terreno, columna_malezas)
            else:
                malezas_terreno = ["Todos"]

            def opciones_con_placeholder(opciones):
                opciones_limpias = [
                    opcion for opcion in opciones
                    if opcion != "Seleccionar"
                ]

                return ["Seleccionar"] + opciones_limpias

            cultivos_terreno = opciones_con_placeholder(cultivos_terreno)
            plagas_terreno = opciones_con_placeholder(plagas_terreno)
            enfermedades_terreno = opciones_con_placeholder(enfermedades_terreno)
            malezas_terreno = opciones_con_placeholder(malezas_terreno)

            valores_modo_terreno = {
                "modo_terreno_cultivo": cultivos_terreno,
                "modo_terreno_plaga": plagas_terreno,
                "modo_terreno_enfermedad": enfermedades_terreno,
                "modo_terreno_maleza": malezas_terreno
            }

            if not st.session_state.get("modo_terreno_placeholder_inicializado", False):
                for clave in valores_modo_terreno:
                    st.session_state[clave] = "Seleccionar"

                st.session_state["modo_terreno_busqueda"] = ""
                st.session_state["modo_terreno_placeholder_inicializado"] = True

            for clave, opciones in valores_modo_terreno.items():
                if st.session_state.get(clave, "Seleccionar") not in opciones:
                    st.session_state[clave] = "Seleccionar"

            col_mt1, col_mt2, col_mt3, col_mt4 = st.columns(4)

            with col_mt1:
                with st.container(border=True):
                    tarjeta_selector_modo_terreno("Cultivo", "🌱")
                    cultivo_terreno = st.selectbox(
                        "Cultivo",
                        cultivos_terreno,
                        key="modo_terreno_cultivo",
                        label_visibility="collapsed"
                    )

            with col_mt2:
                with st.container(border=True):
                    tarjeta_selector_modo_terreno("Plaga", "🐛")
                    plaga_terreno = st.selectbox(
                        "Plaga",
                        plagas_terreno,
                        key="modo_terreno_plaga",
                        label_visibility="collapsed"
                    )

            with col_mt3:
                with st.container(border=True):
                    tarjeta_selector_modo_terreno("Enfermedad", "🦠")
                    enfermedad_terreno = st.selectbox(
                        "Enfermedad",
                        enfermedades_terreno,
                        key="modo_terreno_enfermedad",
                        label_visibility="collapsed"
                    )

            with col_mt4:
                with st.container(border=True):
                    tarjeta_selector_modo_terreno("Maleza", "🌿")
                    maleza_terreno = st.selectbox(
                        "Maleza",
                        malezas_terreno,
                        key="modo_terreno_maleza",
                        label_visibility="collapsed"
                    )

            if not columna_malezas:
                st.caption(
                    "El filtro de malezas estará disponible cuando se incorporen esos datos."
                )

            busqueda_terreno = st.text_input(
                "Buscar producto o dato de etiqueta",
                key="modo_terreno_busqueda",
                placeholder="Ej.: Pirimor, tomate, pulgón..."
            )

            col_buscar, col_limpiar = st.columns([1, 1])

            with col_buscar:
                st.markdown(
                    '<div class="modo-terreno-buscar-marker"></div>',
                    unsafe_allow_html=True
                )
                st.button(
                    "🔎 Buscar",
                    key="modo_terreno_boton_buscar",
                    on_click=buscar_productos_modo_terreno,
                    use_container_width=True
                )

            with col_limpiar:
                st.markdown(
                    '<div class="modo-terreno-limpiar-marker"></div>',
                    unsafe_allow_html=True
                )
                st.button(
                    "↺ Limpiar",
                    key="modo_terreno_eliminar_filtro_inferior",
                    on_click=limpiar_filtros_modo_terreno,
                    use_container_width=True
                )

            criterios_aplicados = st.session_state.get(
                "modo_terreno_criterios_aplicados",
                {}
            )

            busqueda_ejecutada = st.session_state.get(
                "modo_terreno_busqueda_ejecutada",
                False
            )

            hay_consulta_terreno = busqueda_ejecutada and bool(criterios_aplicados)

            if not hay_consulta_terreno:
                recuadro_info_modo_terreno()
            else:
                cultivo_aplicado = criterios_aplicados.get("cultivo", "Seleccionar")
                plaga_aplicada = criterios_aplicados.get("plaga", "Seleccionar")
                enfermedad_aplicada = criterios_aplicados.get(
                    "enfermedad",
                    "Seleccionar"
                )
                maleza_aplicada = criterios_aplicados.get("maleza", "Seleccionar")
                busqueda_aplicada = criterios_aplicados.get("busqueda", "")

                cultivo_filtro = (
                    "Todos" if cultivo_aplicado == "Seleccionar"
                    else cultivo_aplicado
                )
                plaga_filtro = (
                    "Todos" if plaga_aplicada == "Seleccionar"
                    else plaga_aplicada
                )
                enfermedad_filtro = (
                    "Todos" if enfermedad_aplicada == "Seleccionar"
                    else enfermedad_aplicada
                )
                maleza_filtro = (
                    "Todos" if maleza_aplicada == "Seleccionar"
                    else maleza_aplicada
                )

                df_recomendados = filtrar_modo_terreno(
                    df_terreno,
                    cultivo_filtro,
                    plaga_filtro,
                    enfermedad_filtro,
                    maleza_filtro,
                    busqueda_aplicada,
                    columna_malezas
                )

                st.caption(
                    f"{len(df_recomendados)} producto(s) recomendado(s) según los filtros seleccionados."
                )

                if df_recomendados.empty:
                    st.warning(
                        "No se encontraron productos con esos criterios. "
                        "Prueba limpiando un filtro o buscando por nombre comercial."
                    )
                else:
                    st.subheader("Productos recomendados")

                    recuadro_info_modo_terreno()

                    producto_terreno_id_seleccionado = lista_compacta_modo_terreno(
                        df_recomendados,
                        columna_malezas,
                        limite=12
                    )

                    fila_ficha_terreno = None

                    if limpiar_valor(producto_terreno_id_seleccionado) != "":
                        try:
                            producto_terreno_id = int(producto_terreno_id_seleccionado)
                            df_ficha_terreno = df_recomendados[
                                df_recomendados["id"] == producto_terreno_id
                            ]

                            if not df_ficha_terreno.empty:
                                st.session_state["modo_terreno_producto_id"] = str(
                                    producto_terreno_id
                                )
                                fila_ficha_terreno = df_ficha_terreno.iloc[0]
                        except ValueError:
                            fila_ficha_terreno = None

                    if fila_ficha_terreno is not None:
                        ficha_didactica_modo_terreno(
                            fila_ficha_terreno,
                            columna_malezas
                        )

                        pdf_terreno = limpiar_valor(
                            fila_ficha_terreno.get("pdf", "")
                        )

                        try:
                            pdf_terreno_bytes, nombre_pdf_terreno = (
                                obtener_pdf_bytes(pdf_terreno)
                            )
                        except Exception as error:
                            pdf_terreno_bytes = None
                            nombre_pdf_terreno = pdf_terreno
                            st.warning(
                                f"No fue posible abrir el PDF: {error}"
                            )

                        if pdf_terreno_bytes:
                            col_pdf_zoom, col_pdf_descarga = st.columns(2)

                            with col_pdf_zoom:
                                if str(pdf_terreno).startswith(("http://", "https://")):
                                    st.link_button(
                                        "🔍 Abrir PDF con zoom",
                                        pdf_terreno,
                                        use_container_width=True
                                    )
                                else:
                                    st.button(
                                        "🔍 Abrir PDF con zoom",
                                        disabled=True,
                                        use_container_width=True,
                                        help="Disponible para PDF guardados en Supabase."
                                    )

                            with col_pdf_descarga:
                                st.download_button(
                                    label="⬇️ Descargar PDF",
                                    data=pdf_terreno_bytes,
                                    file_name=nombre_pdf_terreno,
                                    mime="application/pdf",
                                    key=f"modo_terreno_pdf_{fila_ficha_terreno['id']}",
                                    use_container_width=True
                                )

                            with st.expander("Ver PDF dentro de la app"):
                                mostrar_pdf_bytes(
                                    pdf_terreno_bytes,
                                    alto=480
                                )
                        else:
                            st.caption("PDF de respaldo no disponible.")

                with st.expander(
                    "🧪 Revisar compatibilidad de mezcla",
                    expanded=False
                ):
                    opciones_compat_terreno = {}

                    for _, fila in df_terreno.iterrows():
                        nombre_compat = limpiar_valor(
                            fila.get("nombre", "")
                        )
                        tipo_compat = limpiar_valor(
                            fila.get("tipo", "")
                        )
                        producto_id_compat = limpiar_valor(
                            fila.get("id", "")
                        )

                        etiqueta = (
                            f"{nombre_compat} | {tipo_compat} | "
                            f"ID {producto_id_compat}"
                        )

                        opciones_compat_terreno[etiqueta] = fila

                    st.caption(
                        "Selecciona 2 a 4 productos para revisar si "
                        "conviene mezclarlos."
                    )

                    opciones_selector_compat = [
                        "Seleccionar producto"
                    ] + list(opciones_compat_terreno.keys())

                    compat_producto_1 = st.selectbox(
                        "Producto 1",
                        opciones_selector_compat,
                        key="compat_producto_1"
                    )

                    compat_producto_2 = st.selectbox(
                        "Producto 2",
                        opciones_selector_compat,
                        key="compat_producto_2"
                    )

                    with st.expander(
                        "Agregar producto 3 y 4 opcional",
                        expanded=False
                    ):
                        compat_producto_3 = st.selectbox(
                            "Producto 3",
                            opciones_selector_compat,
                            key="compat_producto_3"
                        )

                        compat_producto_4 = st.selectbox(
                            "Producto 4",
                            opciones_selector_compat,
                            key="compat_producto_4"
                        )

                    productos_compatibles_labels = [
                        etiqueta
                        for etiqueta in [
                            compat_producto_1,
                            compat_producto_2,
                            compat_producto_3,
                            compat_producto_4
                        ]
                        if etiqueta != "Seleccionar producto"
                    ]

                    productos_compatibles_labels = list(
                        dict.fromkeys(productos_compatibles_labels)
                    )

                    if len(productos_compatibles_labels) < 2:
                        st.info(
                            "Selecciona al menos 2 productos diferentes."
                        )
                    else:
                        compatibilidad_slots_modo_terreno(
                            productos_compatibles_labels
                        )

                        filas_compatibilidad = [
                            opciones_compat_terreno[etiqueta]
                            for etiqueta in productos_compatibles_labels
                        ]

                        try:
                            resultados_compatibilidad = (
                                evaluar_compatibilidad_seleccion(
                                    filas_compatibilidad
                                )
                            )

                            compatibilidad_resultado_modo_terreno(
                                productos_compatibles_labels,
                                resultados_compatibilidad
                            )

                        except Exception as error:
                            st.error(
                                "No fue posible evaluar la compatibilidad: "
                                f"{error}"
                            )

            st.divider()
            botones_finales_modo_terreno()

            accion_modo_terreno = st.query_params.get("accion_modo_terreno", "")

            if accion_modo_terreno == "cargar_pdf":
                formulario_cargar_pdf_modo_terreno()

            elif accion_modo_terreno == "experiencia":
                formulario_experiencia_campo_modo_terreno(df_terreno)


if False:
    # Pantalla antigua Cargar PDF desactivada para dejar solo dos pantallas visibles.
    with tab1:
        st.header("Cargar Etiqueta SAG")

        archivos = st.file_uploader(
            "Selecciona PDFs",
            type=["pdf"],
            accept_multiple_files=True
        )

        if archivos:
            os.makedirs("pdfs", exist_ok=True)

            st.success(f"{len(archivos)} PDF(s) cargado(s)")

            for archivo in archivos:
                st.divider()
                st.subheader(f"Archivo: {archivo.name}")

                ruta = os.path.join("pdfs", archivo.name)

                pdf_bytes = archivo.getvalue()

                with open(ruta, "wb") as f:
                    f.write(pdf_bytes)

                with st.expander("Ver PDF cargado en pantalla"):
                    mostrar_pdf_bytes(pdf_bytes)

                st.download_button(
                    label=f"Descargar PDF: {archivo.name}",
                    data=pdf_bytes,
                    file_name=archivo.name,
                    mime="application/pdf",
                    key=f"descargar_pdf_{archivo.name}"
                )

                texto = extraer_texto_pdf(ruta)

                texto = extraer_texto_pdf(ruta)
                resultado = analizar_texto(texto, archivo.name)

                col1, col2 = st.columns(2)

                with col1:
                    st.write("Producto:", resultado["producto"])
                    st.write("Ingrediente:", resultado["ingrediente"])
                    st.write("Grupo:", resultado["grupo"])
                    st.write("Tipo:", resultado.get("tipo", ""))

                with col2:
                    st.write("Cultivos:", resultado["cultivos"])
                    st.write("Enfermedades:", resultado["enfermedades"])
                    st.write("Insectos:", resultado["insectos"])
                    st.write("Dosis:", resultado.get("dosis", ""))
                    st.write("Compatibilidad:", resultado.get("compatibilidad", ""))
                    st.write("Incompatibilidad:", resultado.get("incompatibilidad", ""))
                    st.write("Fitotoxicidad:", resultado.get("fitotoxicidad", ""))

                st.write("Reingreso:", resultado["reingreso"])
                st.write("Carencia:", resultado["carencia"])
                st.write("Toxicidad abejas:", resultado["toxicidad_abejas"])

                if st.button(f"Guardar {archivo.name} en Base de Datos"):
                    guardar_producto(
                        nombre=resultado["producto"],
                        ingrediente=resultado["ingrediente"],
                        grupo=resultado["grupo"],
                        tipo=resultado.get("tipo", ""),
                        cultivos=resultado["cultivos"],
                        enfermedades=resultado["enfermedades"],
                        insectos=resultado["insectos"],
                        dosis=resultado.get("dosis", ""),
                        compatibilidad=resultado.get("compatibilidad", ""),
                        incompatibilidad=resultado.get("incompatibilidad", ""),
                        fitotoxicidad=resultado.get("fitotoxicidad", ""),
                        reingreso=resultado["reingreso"],
                        carencia=resultado["carencia"],
                        toxicidad_abejas=resultado["toxicidad_abejas"],
                        pdf=archivo.name
                    )

                    st.success(f"{archivo.name} guardado correctamente")

                with st.expander("Ver texto extraido"):
                    st.text_area(
                        f"Texto extraido de {archivo.name}",
                        texto[:5000],
                        height=250
                    )


with tab_base:
    if st.query_params.get(
        "accion_modo_terreno",
        ""
    ) == "base_datos":
        if st.button(
            "← Volver a Modo terreno",
            key="volver_desde_base_datos"
        ):
            st.query_params.clear()
            st.rerun()

    df = obtener_productos()

    columnas_necesarias = [
        "tipo",
        "ingrediente",
        "grupo",
        "carencia",
        "pdf"
    ]

    for columna in columnas_necesarias:
        if columna not in df.columns:
            df[columna] = ""

    if not df.empty:
        df["tipo"] = df.apply(
            completar_tipo_fila,
            axis=1
        )

    # Encabezado y tabla visual.
    base_datos_mockup_header()
    tabla_base_datos_mockup(df)

    st.markdown(
        """
        <style>
        /* Tarjetas inferiores de Base de datos */
        .base-funcional-marker {
            display: none;
        }

        div[data-testid="stHorizontalBlock"]:has(
            .base-funcional-marker
        ) {
            max-width: 920px;
            margin: 18px auto 0 auto;
            gap: 18px;
            align-items: stretch;
        }

        div[data-testid="stVerticalBlock"]:has(
            .base-selector-marker
        ),
        div[data-testid="stVerticalBlock"]:has(
            .base-botones-marker
        ) {
            border: 1px solid #d6dee8;
            border-radius: 18px;
            background: #ffffff;
            padding: 18px;
            box-shadow: 0 9px 26px rgba(15, 23, 42, 0.08);
            min-height: 190px;
        }

        .base-selector-title {
            display: flex;
            align-items: center;
            gap: 12px;
            margin-bottom: 2px;
            color: #0f172a;
            font-size: 1.18rem;
            font-weight: 850;
        }

        .base-selector-icon {
            width: 34px;
            height: 34px;
            border-radius: 10px;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            background: #eafaf2;
            color: #14965a;
            border: 1px solid #bcebd2;
            font-size: 1.1rem;
        }

        .base-selector-subtitle {
            color: #64748b;
            font-size: 0.82rem;
            font-weight: 600;
            margin-left: 46px;
            margin-bottom: 14px;
        }

        div[data-testid="stVerticalBlock"]:has(
            .base-selector-marker
        ) div[data-baseweb="select"] > div {
            border-radius: 12px !important;
            border: 1.5px solid #f0bf54 !important;
            background: #fff9e9 !important;
            min-height: 48px;
        }

        div[data-testid="stVerticalBlock"]:has(
            .base-botones-marker
        ) div.stButton > button,
        div[data-testid="stVerticalBlock"]:has(
            .base-botones-marker
        ) div.stDownloadButton > button {
            width: 100%;
            min-height: 68px;
            border: none !important;
            border-radius: 13px !important;
            background: linear-gradient(
                135deg,
                #0763ca,
                #0753b7
            ) !important;
            color: #ffffff !important;
            font-weight: 850 !important;
            font-size: 0.96rem !important;
            box-shadow: 0 7px 16px rgba(7, 83, 183, 0.22);
            margin-bottom: 8px;
        }

        div[data-testid="stVerticalBlock"]:has(
            .base-botones-marker
        ) div.stButton > button:hover,
        div[data-testid="stVerticalBlock"]:has(
            .base-botones-marker
        ) div.stDownloadButton > button:hover {
            background: linear-gradient(
                135deg,
                #0757b5,
                #06469a
            ) !important;
            transform: translateY(-1px);
        }

        .base-producto-seleccionado {
            max-width: 920px;
            margin: 10px auto 0 auto;
            color: #64748b;
            font-size: 0.84rem;
        }

        .base-admin-wrapper {
            max-width: 920px;
            margin: 12px auto 0 auto;
        }

        @media screen and (max-width: 700px) {
            div[data-testid="stHorizontalBlock"]:has(
                .base-funcional-marker
            ) {
                flex-direction: column !important;
                gap: 10px !important;
            }

            div[data-testid="stVerticalBlock"]:has(
                .base-selector-marker
            ),
            div[data-testid="stVerticalBlock"]:has(
                .base-botones-marker
            ) {
                padding: 14px !important;
                min-height: auto !important;
                border-radius: 15px !important;
            }

            div[data-testid="stVerticalBlock"]:has(
                .base-botones-marker
            ) div.stButton > button,
            div[data-testid="stVerticalBlock"]:has(
                .base-botones-marker
            ) div.stDownloadButton > button {
                min-height: 54px !important;
                font-size: 0.88rem !important;
            }
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    if df.empty:
        st.info("Aún no hay productos guardados.")

    else:
        opciones_productos = {}

        for _, fila in df.iterrows():
            nombre = str(
                fila.get("nombre", "")
            ).strip()

            tipo = str(
                fila.get("tipo", "")
            ).strip()

            producto_id = int(fila["id"])

            etiqueta = (
                f"{nombre} | {tipo} | ID {producto_id}"
            )

            opciones_productos[etiqueta] = fila

        col_selector, col_botones = st.columns(
            [1.15, 0.75]
        )

        with col_selector:
            st.markdown(
                '<div class="base-funcional-marker base-selector-marker"></div>',
                unsafe_allow_html=True
            )

            st.markdown(
                """
                <div class="base-selector-title">
                    <span class="base-selector-icon">▯</span>
                    <span>Seleccionar producto</span>
                </div>
                <div class="base-selector-subtitle">
                    Busca y selecciona un producto de la base de datos.
                </div>
                """,
                unsafe_allow_html=True
            )

            producto_seleccionado = st.selectbox(
                "Producto",
                list(opciones_productos.keys()),
                key="base_producto_seleccionado",
                label_visibility="collapsed"
            )

        fila_seleccionada = opciones_productos[
            producto_seleccionado
        ]

        producto_id = int(
            fila_seleccionada["id"]
        )

        nombre_producto = str(
            fila_seleccionada.get(
                "nombre",
                ""
            )
        ).strip()

        referencia_pdf = str(
            fila_seleccionada.get(
                "pdf",
                ""
            )
        ).strip()

        pdf_bytes = None
        nombre_pdf = ""

        if (
            referencia_pdf
            and referencia_pdf.lower() != "nan"
        ):
            try:
                pdf_bytes, nombre_pdf = obtener_pdf_bytes(
                    referencia_pdf
                )
            except Exception as error:
                st.warning(
                    "No fue posible obtener el PDF: "
                    f"{error}"
                )

        with col_botones:
            st.markdown(
                '<div class="base-funcional-marker base-botones-marker"></div>',
                unsafe_allow_html=True
            )

            if st.button(
                "▣  Ver PDF del producto seleccionado",
                key=f"base_ver_pdf_{producto_id}",
                use_container_width=True,
                disabled=not bool(pdf_bytes)
            ):
                st.session_state[
                    "base_pdf_visible_id"
                ] = producto_id

            if pdf_bytes:
                st.download_button(
                    "▣  Descargar PDF seleccionado",
                    data=pdf_bytes,
                    file_name=(
                        nombre_pdf
                        or f"{nombre_producto}.pdf"
                    ),
                    mime="application/pdf",
                    key=(
                        "base_descargar_pdf_"
                        f"{producto_id}"
                    ),
                    use_container_width=True
                )
            else:
                st.button(
                    "▣  Descargar PDF seleccionado",
                    key=(
                        "base_descargar_pdf_"
                        f"desactivado_{producto_id}"
                    ),
                    disabled=True,
                    use_container_width=True
                )

        st.markdown(
            (
                '<div class="base-producto-seleccionado">'
                f'Producto seleccionado: '
                f'<strong>{html.escape(nombre_producto)}</strong>'
                '</div>'
            ),
            unsafe_allow_html=True
        )

        if (
            st.session_state.get(
                "base_pdf_visible_id"
            ) == producto_id
            and pdf_bytes
        ):
            with st.expander(
                f"PDF de {nombre_producto}",
                expanded=True
            ):
                mostrar_pdf_bytes(
                    pdf_bytes,
                    alto=650
                )

        st.markdown(
            '<div class="base-admin-wrapper">',
            unsafe_allow_html=True
        )

        with st.expander(
            "⚙️ Administración avanzada",
            expanded=False
        ):
            st.caption(
                "Esta sección permite eliminar un producto "
                "y su archivo PDF."
            )

            with st.expander(
                "🗑️ Eliminar producto",
                expanded=False
            ):
                st.warning(
                    "Esta acción es permanente. Se eliminará "
                    "el producto de la base de datos y también "
                    "su PDF almacenado."
                )

                confirmar = st.checkbox(
                    (
                        "Confirmo que deseo eliminar "
                        f"{nombre_producto}"
                    ),
                    key=(
                        "base_confirmar_eliminar_"
                        f"{producto_id}"
                    )
                )

                if st.button(
                    "Eliminar producto y PDF",
                    key=f"base_eliminar_{producto_id}",
                    use_container_width=True,
                    disabled=not confirmar
                ):
                    try:
                        with st.spinner(
                            "Eliminando producto y PDF..."
                        ):
                            if (
                                referencia_pdf
                                and referencia_pdf.lower()
                                != "nan"
                            ):
                                eliminar_pdf_storage(
                                    referencia_pdf
                                )

                            eliminar_producto(
                                producto_id
                            )

                        st.session_state.pop(
                            "base_pdf_visible_id",
                            None
                        )

                        st.success(
                            f"{nombre_producto} fue eliminado."
                        )

                        st.rerun()

                    except Exception as error:
                        st.error(
                            "No fue posible eliminar el producto: "
                            f"{error}"
                        )

        st.markdown(
            "</div>",
            unsafe_allow_html=True
        )

if False:
    # Pantalla antigua Manual de uso desactivada para dejar solo dos pantallas visibles.
    with tab3:
        st.header("Manual de uso - IA Agrícola V15")

        st.subheader("1. Cargar PDF")
        st.write(
            "En la pestaña Cargar PDF puedes subir una o varias etiquetas SAG en formato PDF. "
            "La app extrae automáticamente información como producto, ingrediente activo, tipo, dosis, cultivos, "
            "problemas, carencia, reingreso, toxicidad para abejas y compatibilidad cuando la etiqueta lo indique."
        )

        st.subheader("2. Guardar producto")
        st.write(
            "Después de cargar un PDF, revisa los datos extraídos. Si están correctos, presiona el botón "
            "Guardar en Base de Datos. El producto quedará disponible en la pestaña Base de datos."
        )

        st.subheader("3. Base de datos")
        st.write(
            "En Base de datos puedes revisar todos los productos guardados, filtrar por tipo, cultivo, ingrediente "
            "activo o problema, descargar la base en Excel y eliminar productos repetidos o seleccionados."
        )

        st.subheader("4. Vista visual del producto")
        st.write(
            "La vista visual muestra la información principal del producto seleccionado en formato de ficha rápida: "
            "grupo químico, composición, dosis, cultivos, problema, carencia, reingreso, abejas y compatibilidad."
        )

        st.subheader("5. Compatibilidad")
        st.write(
            "El comparador permite seleccionar dos productos y revisar si alguna etiqueta menciona compatibilidad "
            "o incompatibilidad entre ellos."
        )

        st.warning(
            "Importante: si la app indica 'No hay información suficiente', no significa que los productos sean incompatibles. "
            "Significa que en las etiquetas guardadas no se encontró una mención clara entre ambos productos."
        )

        st.subheader("6. Compatibilidad manual")
        st.write(
            "Puedes completar o corregir compatibilidad, incompatibilidad y fitotoxicidad manualmente. "
            "Esto es útil cuando la etiqueta no se leyó correctamente o cuando deseas registrar información revisada."
        )

        st.subheader("7. Ficha técnica Word")
        st.write(
            "La app permite descargar una ficha técnica en Word del producto seleccionado. "
            "La ficha incluye datos técnicos, compatibilidad, incompatibilidad, fitotoxicidad y advertencia de revisión SAG."
        )

        st.subheader("8. PDF de respaldo")
        st.write(
            "Desde Base de datos puedes ver el PDF original asociado al producto y también descargarlo."
        )

        st.subheader("9. Respaldo completo ZIP")
        st.write(
            "El botón Descargar respaldo completo ZIP genera una copia con la base agro.db, un Excel de productos "
            "y los PDF guardados. Es recomendable descargar este respaldo cuando avances bastante en tu base."
        )

        st.subheader("10. Recomendación de uso")
        st.error(
            "Antes de aplicar o mezclar productos, revisar siempre la etiqueta SAG oficial, dosis autorizada, "
            "cultivo autorizado, carencia, reingreso, compatibilidad e incompatibilidades."
        )

# Visor PDF compatible con Streamlit Cloud
def mostrar_pdf_bytes(pdf_bytes, alto=850):
    import hashlib
    import fitz
    import streamlit as st

    if not pdf_bytes:
        st.warning("No se pudo cargar el PDF.")
        return

    try:
        documento = fitz.open(
            stream=pdf_bytes,
            filetype="pdf"
        )

        total_paginas = documento.page_count

        if total_paginas == 0:
            st.warning("El PDF no contiene páginas.")
            documento.close()
            return

        identificador = hashlib.sha1(pdf_bytes).hexdigest()[:12]

        pagina_seleccionada = st.number_input(
            "Página del PDF",
            min_value=1,
            max_value=total_paginas,
            value=1,
            step=1,
            key=f"pagina_pdf_{identificador}"
        )

        pagina = documento.load_page(
            int(pagina_seleccionada) - 1
        )

        matriz = fitz.Matrix(1.7, 1.7)
        imagen = pagina.get_pixmap(
            matrix=matriz,
            alpha=False
        ).tobytes("png")

        st.caption(
            f"Página {int(pagina_seleccionada)} de {total_paginas}"
        )

        st.image(
            imagen,
            use_container_width=True
        )

        documento.close()

    except Exception as error:
        st.error(f"No fue posible visualizar el PDF: {error}")

